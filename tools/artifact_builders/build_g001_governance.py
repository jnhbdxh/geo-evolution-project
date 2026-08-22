from __future__ import annotations

import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(Path(__file__).resolve().parent))

from build_arch_reconciliation_v11 import (  # noqa: E402
    AMBER,
    BLUE,
    DARK_BLUE,
    GREEN,
    INK,
    MUTED,
    add_bullet,
    add_callout,
    add_field,
    add_numbered,
    add_para,
    add_table,
    configure_styles,
    h1,
    h2,
    set_run_font,
)


DEFAULT_OUTPUT = ROOT / "outputs" / "g0-01" / "GEO_OS_G0-01_Authority_Change_Control_ADR-001_V1.0.docx"
OUTPUT = Path(os.environ.get("GEO_OS_ARTIFACT_OUTPUT", DEFAULT_OUTPUT))


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("GEO OS  |  G0-01 Scope, Authority & Change Control")
    set_run_font(r, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    r = fp.add_run("GEO OS · G0-01 Review Pack  |  Page ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(fp, "PAGE")
    r = fp.add_run(" of ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(fp, "NUMPAGES")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("GEO OS")
    set_run_font(r, size=12, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("G0-01 Scope, Authority & Change Control Pack")
    set_run_font(r, size=23, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Scope Matrix · ADR-001 · Change Control · Open Decision Register")
    set_run_font(r, size=14, color=DARK_BLUE)

    add_callout(
        doc,
        "当前结论",
        "G0-01 保持 OPEN / PENDING REVIEW。本文件把退出条件转为可评审契约，但不代替正式批准，也不提前关闭 Gate。",
        fill=AMBER,
        accent="7A5A00",
    )

    add_table(
        doc,
        ["项目", "当前口径"],
        [
            ["文档版本", "V1.0"],
            ["基准日期", "2026-08-21"],
            ["ADR", "ADR-001 — Document Authority by Decision Domain（PROPOSED）"],
            ["配套矩阵", "GEO_OS_G0-01_Scope_Matrix_V1.0.xlsx"],
            ["范围项", "39 项：32 IN / 2 OUT / 5 DEFERRED"],
            ["未决事项", "6 项 Open Decision；39 项具名负责人待指定"],
            ["Gate 状态", "OPEN / PENDING REVIEW"],
        ],
        [2160, 7200],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("关闭 G0-01 必须以评审证据为基础：范围完整、权威明确、冲突可处理、变更可追踪、负责人和期限可执行。")
    set_run_font(r, size=10.5, color=MUTED, italic=True)
    doc.add_page_break()


def build() -> None:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_cover(doc)

    h1(doc, "1. 执行结论与退出条件")
    add_para(
        doc,
        "G0-01 的任务不是证明所有未来需求都已设计完成，而是证明当前开发范围、文档权威、冲突处理和变更治理已经足以防止跨模块返工。",
    )
    add_callout(
        doc,
        "Gate 状态",
        "OPEN / PENDING REVIEW。Scope Matrix 已形成初稿，ADR-001 与 Change Control 已形成待批准版本；具名 Owner、Open Decision 和正式批准记录尚未关闭。",
        fill=AMBER,
        accent="7A5A00",
    )
    add_table(
        doc,
        ["退出条件", "要求", "当前状态"],
        [
            ["范围覆盖", "覆盖所有模块、运行单元和跨模块能力；逐项明确 IN / OUT / DEFERRED、MVP 等级、权威来源、验收与负责人。", "DRAFT COMPLETE"],
            ["领域化权威", "按商业方向、MVP 范围与验收、领域语义、技术结构、工程实现划分。", "ADR PROPOSED"],
            ["冲突规则", "同一领域内采用正式冻结、范围更具体、显式替代且版本更新的契约。", "ADR PROPOSED"],
            ["变更流程", "Change Request 至 Release 的七阶段闭环；四类变更有不同批准与版本要求。", "DRAFTED"],
            ["已知冲突", "全部解决，或登记具备负责人、期限和关闭产物的 Open Decision。", "6 OPEN"],
            ["正式关闭", "记录批准人、ADR、版本和关闭时间。", "NOT APPROVED"],
        ],
        [1760, 5900, 1700],
        compact=True,
    )

    h1(doc, "2. Scope Matrix 契约")
    add_para(
        doc,
        "配套工作簿是 G0-01 的范围事实表。任何模块、跨模块能力或运行单元在进入正式开发前，必须能够定位到稳定 Scope ID。",
    )
    add_table(
        doc,
        ["字段组", "强制字段", "治理规则"],
        [
            ["范围决策", "Scope ID、Type、Domain、Capability、IN / OUT / DEFERRED", "决策值只允许三态；任何变化属于受控变更。"],
            ["MVP 与交付", "MVP Level、Delivery Phase、Gate / Dependency", "区分基础能力、Walking Skeleton、Commercial Complete 与 Post-MVP。"],
            ["权威与实现", "Authority Domain、Primary Authority、Implementation Authority", "不得以工程实现反向修改上位商业、范围或领域语义。"],
            ["可验收性", "Acceptance Criteria、Known Limitation", "必须是可测试结果，而不是原则描述。"],
            ["责任", "Owner Role、Named Owner", "当前 Role Owner 已填；具名 Owner 必须在关闭前指定。"],
        ],
        [1700, 3440, 4220],
    )
    h2(doc, "2.1 当前矩阵摘要")
    add_table(
        doc,
        ["分类", "数量", "说明"],
        [
            ["IN", "32", "进入 V1 的基础、Walking Skeleton 或 Commercial Complete 范围。"],
            ["OUT", "2", "第三方高级平台能力与动态低代码规则平台不在 V1。"],
            ["DEFERRED", "5", "需后续语义或范围决议，但不阻塞非争议基础设施。"],
            ["总计", "39", "15 个模块、4 个运行单元及 20 项跨模块能力。"],
        ],
        [1500, 900, 6960],
    )
    add_bullet(doc, "Scope Decision 为空的行数必须为 0。")
    add_bullet(doc, "每行具名 Owner 必须完成；Role Owner 不能替代最终责任人。")
    add_bullet(doc, "OUT 或 DEFERRED 的能力若提前进入开发，必须发起 MVP_SCOPE_CHANGE。")

    h1(doc, "3. ADR-001 — Document Authority by Decision Domain")
    add_table(
        doc,
        ["ADR 字段", "值"],
        [
            ["Status", "PROPOSED"],
            ["Decision Owner", "Architecture Governance Lead（具名负责人待指定）"],
            ["Decision Date", "2026-08-21"],
            ["Gate", "G0-01"],
            ["Supersedes", "任何简单全局文档排名口径"],
        ],
        [1900, 7460],
    )
    h2(doc, "3.1 Context")
    add_para(
        doc,
        "业务背景、产品、需求、领域决议、技术方案和工程 Contract 同时存在。它们解决的决策问题不同，不能用一条简单的全局排名解释全部冲突。",
    )
    h2(doc, "3.2 Decision")
    add_table(
        doc,
        ["决策领域", "首要权威", "控制内容", "工程输出"],
        [
            ["商业方向", "业务背景 / 产品方案", "商业目标、目标客户、产品边界、商业模式、路线方向", "Scope Change Request / Product Decision"],
            ["MVP 范围与验收", "需求方案", "IN/OUT、优先级、角色权限、功能与非功能验收", "Scope Matrix / Acceptance ID"],
            ["领域语义", "A1/A2/A3/B/C 最终冻结决议包", "Observation、Eligibility、Review、Mention、Recommendation、Citation、Snapshot 等", "Domain Contract / Semantic ADR"],
            ["技术结构", "Technical Architecture Reconciliation / 技术方案", "模块边界、运行单元、存储、事件、可靠性、安全和基础设施", "Architecture ADR"],
            ["工程实现", "Domain Contract、ADR、API/Schema/Event、Migration、Contract Test", "可执行实现与版本", "Code / Test / Release Record"],
        ],
        [1450, 2400, 3520, 1990],
        compact=True,
    )
    h2(doc, "3.3 同一决策领域内的冲突处理")
    add_numbered(doc, "先识别冲突属于哪个决策领域；不得跨领域直接比较文档级别。")
    add_numbered(doc, "优先采用已正式冻结或已批准的契约，而不是草案、讨论稿或口头结论。")
    add_numbered(doc, "在冻结状态相同的情况下，范围更具体的契约优先于一般性描述。")
    add_numbered(doc, "版本更新只有在显式声明替代关系时才优先；不得用“文件日期较新”静默覆盖旧语义。")
    add_numbered(doc, "跨领域冲突或无法消解的冲突必须进入 Change Request 和 Open Decision，普通开发 PR 不得自行裁决。")
    h2(doc, "3.4 Consequences")
    add_bullet(doc, "工程实现必须证明其与 Scope、领域语义和技术结构的追踪关系。")
    add_bullet(doc, "需求方案不能修改 A/B/C 语义；A/B/C 也不直接决定商业发布时间和计费边界。")
    add_bullet(doc, "破坏性契约变更可能触发相关 Gate 重开，而不是只升级代码版本。")

    h1(doc, "4. Change Control")
    add_callout(
        doc,
        "标准流程",
        "Change Request → 分类 → 影响评估 → 审批 → 版本升级 → Traceability 更新 → 发布",
    )
    add_table(
        doc,
        ["阶段", "必须完成", "可执行产物", "责任角色"],
        [
            ["1. Change Request", "说明问题、动机、拟议变化、受影响对象和期望时间。", "Request ID + Evidence", "Requester / PM"],
            ["2. 分类", "确定编辑性、兼容性新增、破坏性契约变更或 MVP 范围变更。", "Change Class", "Architecture Governance"],
            ["3. 影响评估", "评估 Scope、语义、API/Schema/Event、迁移、历史复现、安全、测试和交付。", "Impact Report", "Domain + Technical Owners"],
            ["4. 审批", "按决策领域路由；破坏性和范围变更不得由单一开发者批准。", "Approval Record", "Designated Approver"],
            ["5. 版本升级", "升级文档、Contract、API/Schema/Event 和 Policy Release。", "Versioned Artifacts", "Artifact Owners"],
            ["6. Traceability", "更新 Scope、ADR、Rule→Object→Contract→Code→Test 和 Open Decision。", "Updated Traceability", "Governance + QA"],
            ["7. 发布", "通过迁移、兼容、回归、契约和回滚检查。", "Release Record", "Release Owner"],
        ],
        [1450, 3980, 2250, 1680],
        compact=True,
    )
    h2(doc, "4.1 变更分类")
    add_table(
        doc,
        ["分类", "定义", "版本影响", "批准", "最低证据"],
        [
            ["EDITORIAL", "不改变语义、范围、字段含义或行为的文字/排版修正。", "Patch", "Document Owner", "校对；必要时更新引用。"],
            ["COMPATIBLE_ADDITION", "新增可选字段、状态或不影响旧消费者的能力。", "Minor", "Domain + Technical Owner", "向后兼容、默认行为与契约测试。"],
            ["BREAKING_CONTRACT", "改变既有语义、必填字段、状态机、约束、历史复现或消费者行为。", "Major", "Domain + Architecture Approver", "迁移、兼容/回滚、历史影响和完整契约回归。"],
            ["MVP_SCOPE_CHANGE", "改变 IN/OUT/DEFERRED、MVP 等级、验收、第三方能力或正式 KPI 准入。", "Scope Baseline", "Product Sponsor + PM", "商业、排期、成本、依赖、验收和风险评估。"],
        ],
        [1460, 3020, 1330, 1850, 1700],
        compact=True,
    )
    h2(doc, "4.2 影响评估的强制检查")
    add_bullet(doc, "是否改变 Tenant ownership、授权或跨租户可见性。")
    add_bullet(doc, "是否改变 Raw Fact、Assessment、Resolution、Snapshot 或历史复现。")
    add_bullet(doc, "是否需要数据迁移、双读写、兼容窗口、回滚或消费者升级。")
    add_bullet(doc, "是否改变 Policy Release、算法、Prompt、Metric 或 Coverage Context。")
    add_bullet(doc, "是否更新 Scope Matrix、ADR、API/Schema/Event、Contract Test 和 Release Record。")

    h1(doc, "5. 已知冲突与 Open Decision")
    h2(doc, "5.1 已解决并冻结的冲突")
    add_table(
        doc,
        ["编号", "冲突", "冻结处理"],
        [
            ["RC-01", "产品基线将 Agency/White-label/API 放在 V1 之后，但系统需供第三方使用。", "V1 仅受控开通 Tenant Workspace 并管理自有客户；高级平台代理能力 OUT。"],
            ["RC-02", "Tenant→Customer→Brand→Project 访问链与全局 Source 身份复用。", "采用 GLOBAL_IDENTITY_WITH_TENANT_CONTEXT；全局目录不向租户直接暴露。"],
            ["RC-03", "策略需版本化与是否建设动态规则平台。", "V1 使用 Versioned Manifest + Immutable Release + Code Evaluator；动态低代码平台 OUT。"],
            ["RC-04", "Project Policy 默认配置与历史实际执行版本。", "Execution、Assessment、Resolution、Snapshot 分别保存实际 policy_release_id。"],
            ["RC-05", "Walking Skeleton 报告链与正式客户 KPI。", "只使用测试指标或冻结 ANSWER_OUTCOME；Basic Report 只验证 Snapshot 消费。"],
        ],
        [900, 3710, 4750],
        compact=True,
    )
    h2(doc, "5.2 Open Decision Register")
    add_table(
        doc,
        ["ID", "未决事项", "Owner", "Due", "Status", "关闭产物"],
        [
            ["OD-G01-001", "为 39 项 Scope 指定具名负责人和替补负责人。", "PM / Sponsor", "2026-08-25", "OPEN", "Named Owner columns completed"],
            ["OD-G01-002", "确认 GEO Intelligence 正式策略输出的 Commercial MVP 边界。", "Product Owner", "2026-08-26", "OPEN", "Scope + acceptance update"],
            ["OD-G01-003", "确认 Intervention / Effect Validation 的 Commercial MVP 边界。", "Product + Domain", "2026-08-27", "OPEN", "Intervention acceptance contract"],
            ["OD-G01-004", "确认 Publisher Worker 在 V1 中是 DEFERRED 还是 MVP-2。", "Technical + Product", "2026-08-25", "OPEN", "Runtime scope ADR"],
            ["OD-G01-005", "完成需求章节至 Scope/Acceptance/Test 的覆盖核对。", "BA + QA", "2026-08-28", "OPEN", "Requirements coverage report"],
            ["OD-G01-006", "冻结正式 Mention/Recommendation/Citation KPI 准入条件。", "Measurement + Product", "2026-08-27", "OPEN", "KPI entry ADR + traceability"],
        ],
        [900, 3060, 1450, 1050, 900, 2000],
        compact=True,
    )

    h1(doc, "6. G0-01 关闭评估与记录")
    add_table(
        doc,
        ["检查项", "当前证据", "状态", "关闭条件"],
        [
            ["Scope coverage", "39 项已编目，包含 15 模块、4 运行单元和 20 跨模块能力。", "DRAFT COMPLETE", "评审确认无遗漏。"],
            ["Mandatory fields", "决策、MVP、权威、验收和 Role Owner 已填；Named Owner 未填。", "OPEN", "39 项具名 Owner 全部指定。"],
            ["Authority ADR", "ADR-001 已形成 PROPOSED 版本。", "PENDING REVIEW", "批准版本和批准人记录。"],
            ["Change Control", "七阶段、四分类和影响评估清单已形成。", "PENDING REVIEW", "确认批准角色和执行 SLA。"],
            ["Known conflicts", "5 项已冻结处理；6 项 Open Decision 已登记。", "OPEN", "Open Decision 全部解决或正式接受。"],
            ["Gate record", "尚无关闭批准记录。", "OPEN", "填写批准人、ADR、版本和关闭时间。"],
        ],
        [1900, 3860, 1500, 2100],
        compact=True,
    )
    add_callout(
        doc,
        "评审结论",
        "本轮产物达到“提交 G0-01 评审”的条件，但未达到“关闭 G0-01”的条件。正式状态继续为 OPEN / PENDING REVIEW。",
        fill=GREEN,
        accent="2F6B45",
    )

    h2(doc, "6.1 Gate 关闭记录模板")
    add_table(
        doc,
        ["字段", "填写内容"],
        [
            ["Gate ID", "G0-01"],
            ["Owner", ""],
            ["Decision / ADR", "ADR-001 approved version"],
            ["Executable Artifact", "Scope Matrix + Authority Map + Change Control + Open Decision Register"],
            ["Contract Tests", "Scope completeness / traceability checks"],
            ["Known Limitations", ""],
            ["Status", "OPEN"],
            ["Approved At", ""],
            ["Approved By", ""],
        ],
        [2400, 6960],
    )

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    if OUTPUT.exists():
        raise FileExistsError(
            f"Refusing to overwrite existing artifact: {OUTPUT}. "
            "Set GEO_OS_ARTIFACT_OUTPUT to a new versioned path."
        )
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
