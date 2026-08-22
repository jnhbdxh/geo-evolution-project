from __future__ import annotations

import html
import math
import os
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = Path(
    os.environ.get(
        "GEO_OS_PRODUCT_OUTPUT_DIR",
        str(REPO_ROOT / "outputs" / "product-v1.1"),
    )
).resolve()
DOCX_OUTPUT = OUTPUT_DIR / "GEO_OS_产品方案对齐修订版_V1.1.docx"
BUSINESS_SVG = OUTPUT_DIR / "GEO_OS_业务架构图_V1.1.svg"
BUSINESS_PNG = OUTPUT_DIR / "GEO_OS_业务架构图_V1.1.png"
TECH_SVG = OUTPUT_DIR / "GEO_OS_技术架构图_V1.1.svg"
TECH_PNG = OUTPUT_DIR / "GEO_OS_技术架构图_V1.1.png"

SOURCE_DOCS = [
    r"D:\三合星链\技术组\v2\项目文档\整体设计\GEO_OS_业务背景定稿版_V1.0.docx",
    r"D:\三合星链\技术组\v2\项目文档\整体设计\GEO_OS_产品方案定稿版_V1.0.docx",
    r"D:\三合星链\技术组\v2\项目文档\整体设计\GEO_OS_技术方案定稿版_V1.0.docx",
    r"D:\三合星链\技术组\v2\项目文档\整体设计\GEO_OS_需求方案定稿版_V1.0.docx",
]

FONT_REGULAR = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\msyhbd.ttc")

NAVY = "0B2545"
BLUE = "2E74B5"
TEAL = "167D8D"
GREEN = "2F6B45"
AMBER = "9A6700"
RED = "9B1C1C"
INK = "1F2937"
GRAY = "667085"
LINE = "CBD5E1"
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "E7F4F5"
LIGHT_GREEN = "E8F3EC"
LIGHT_AMBER = "FFF4CE"
LIGHT_RED = "FDECEC"
LIGHT = "F7F9FC"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, *, size: float = 11, bold: bool = False, color: str = INK, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int], indent: int = 120) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths)):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def style_table(table, widths: Sequence[int], *, header_fill: str = LIGHT_BLUE, font_size: float = 9.2) -> None:
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, bold=row_index == 0, color=NAVY if row_index == 0 else INK)


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, text, separate, result, end])
    set_run_font(run, size=8.5, color=GRAY)


def add_rule(paragraph, color: str = BLUE, size: int = 10) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run("GEO OS  |  产品方案对齐修订版 V1.1")
    set_run_font(run, size=8.5, bold=True, color=GRAY)
    add_rule(hp, color=LINE, size=4)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("GEO OS · Product Alignment  |  ")
    set_run_font(run, size=8.5, color=GRAY)
    add_field(fp, "PAGE")
    run = fp.add_run(" / ")
    set_run_font(run, size=8.5, color=GRAY)
    add_field(fp, "NUMPAGES")


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("PRODUCT ALIGNMENT BASELINE")
    set_run_font(run, size=9.5, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("GEO OS（生成式引擎优化操作系统）")
    set_run_font(run, size=25, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("产品方案对齐修订版 V1.1")
    set_run_font(run, size=17, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("依据最新业务规划、Gate 0 推进方式和单负责人治理结论形成")
    set_run_font(run, size=11, color=GRAY)

    metadata = doc.add_table(rows=5, cols=2)
    rows = [
        ("文档状态", "对齐修订基线 / G0-01 OPEN · PENDING DOCUMENTATION COMPLETION"),
        ("版本与日期", "V1.1 / 2026-08-21"),
        ("批准口径", "项目唯一负责人已确认现行业务规划与系统逻辑"),
        ("适用范围", "GEO OS Core Platform V1、Walking Skeleton 与后续语义扩展"),
        ("版本关系", "不覆盖四份 V1.0；本文件记录最新讨论形成的产品修订"),
    ]
    for index, (label, value) in enumerate(rows):
        metadata.cell(index, 0).text = label
        metadata.cell(index, 1).text = value
    style_table(metadata, [1900, 7460], header_fill=LIGHT_BLUE, font_size=9.2)
    for row in metadata.rows:
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=9.2, bold=True, color=NAVY)
    doc.add_paragraph()


def add_callout(doc: Document, label: str, text: str, *, fill: str = LIGHT_TEAL, accent: str = TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    # Treat the single semantic callout row as its accessible header row so
    # assistive technology receives an explicit table-row role.
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=8.8, color=GRAY, italic=True)


def set_picture_alt(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


class DiagramCanvas:
    def __init__(self, width: int, height: int, title: str, subtitle: str) -> None:
        self.width = width
        self.height = height
        self.title = title
        self.subtitle = subtitle
        self.image = Image.new("RGB", (width, height), "#F8FAFC")
        self.draw = ImageDraw.Draw(self.image)
        self.svg: list[str] = [
            f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
            '<defs><marker id="arrow" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto" markerUnits="strokeWidth"><path d="M0,0 L0,6 L9,3 z" fill="#64748B"/></marker></defs>',
            '<rect width="100%" height="100%" fill="#F8FAFC"/>',
        ]
        self.fonts = {
            "title": ImageFont.truetype(str(FONT_BOLD), 42),
            "subtitle": ImageFont.truetype(str(FONT_REGULAR), 22),
            "section": ImageFont.truetype(str(FONT_BOLD), 23),
            "box": ImageFont.truetype(str(FONT_BOLD), 20),
            "body": ImageFont.truetype(str(FONT_REGULAR), 17),
            "small": ImageFont.truetype(str(FONT_REGULAR), 15),
        }
        self.text(60, 38, title, font="title", fill="#0B2545")
        self.text(60, 92, subtitle, font="subtitle", fill="#667085")
        self.line(60, 130, width - 60, 130, fill="#2E74B5", width=4)

    def text(self, x: int, y: int, value: str, *, font: str = "body", fill: str = "#1F2937", anchor: str = "la") -> None:
        self.draw.text((x, y), value, font=self.fonts[font], fill=fill, anchor=anchor)
        family = "Microsoft YaHei"
        size = {"title": 42, "subtitle": 22, "section": 23, "box": 20, "body": 17, "small": 15}[font]
        weight = "700" if font in {"title", "section", "box"} else "400"
        svg_anchor = {"la": "start", "ma": "middle", "ra": "end"}.get(anchor, "start")
        baseline = y + size * 0.85
        self.svg.append(
            f'<text x="{x}" y="{baseline:.1f}" fill="{fill}" font-family="{family}" font-size="{size}" font-weight="{weight}" text-anchor="{svg_anchor}">{html.escape(value)}</text>'
        )

    def wrapped_text(self, x: int, y: int, width: int, value: str, *, font: str = "body", fill: str = "#1F2937", line_gap: int = 7, center: bool = False) -> int:
        words = list(value)
        lines: list[str] = []
        current = ""
        for char in words:
            candidate = current + char
            if self.draw.textlength(candidate, font=self.fonts[font]) <= width or not current:
                current = candidate
            else:
                lines.append(current)
                current = char
        if current:
            lines.append(current)
        size = {"title": 42, "subtitle": 22, "section": 23, "box": 20, "body": 17, "small": 15}[font]
        for index, line in enumerate(lines):
            yy = y + index * (size + line_gap)
            xx = x + width // 2 if center else x
            self.text(xx, yy, line, font=font, fill=fill, anchor="ma" if center else "la")
        return len(lines) * (size + line_gap)

    def rect(self, x: int, y: int, w: int, h: int, *, fill: str, outline: str = "#CBD5E1", radius: int = 16, line_width: int = 2) -> None:
        self.draw.rounded_rectangle((x, y, x + w, y + h), radius=radius, fill=fill, outline=outline, width=line_width)
        self.svg.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{radius}" fill="{fill}" stroke="{outline}" stroke-width="{line_width}"/>')

    def line(self, x1: int, y1: int, x2: int, y2: int, *, fill: str = "#64748B", width: int = 3) -> None:
        self.draw.line((x1, y1, x2, y2), fill=fill, width=width)
        self.svg.append(f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="{fill}" stroke-width="{width}"/>')

    def arrow(self, x1: int, y1: int, x2: int, y2: int, *, fill: str = "#64748B", width: int = 3) -> None:
        self.draw.line((x1, y1, x2, y2), fill=fill, width=width)
        angle = math.atan2(y2 - y1, x2 - x1)
        length = 13
        points = [
            (x2, y2),
            (x2 - length * math.cos(angle - math.pi / 6), y2 - length * math.sin(angle - math.pi / 6)),
            (x2 - length * math.cos(angle + math.pi / 6), y2 - length * math.sin(angle + math.pi / 6)),
        ]
        self.draw.polygon(points, fill=fill)
        self.svg.append(f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="{fill}" stroke-width="{width}" marker-end="url(#arrow)"/>')

    def box(self, x: int, y: int, w: int, h: int, title: str, lines: Sequence[str], *, fill: str = "#FFFFFF", outline: str = "#CBD5E1", title_color: str = "#0B2545") -> None:
        self.rect(x, y, w, h, fill=fill, outline=outline)
        self.text(x + 18, y + 14, title, font="box", fill=title_color)
        yy = y + 48
        for line in lines:
            used_height = self.wrapped_text(x + 18, yy, w - 36, line, font="small", fill="#475467")
            yy += max(24, used_height)

    def save(self, png_path: Path, svg_path: Path) -> None:
        self.svg.append("</svg>")
        self.image.save(png_path, format="PNG", optimize=True)
        svg_path.write_text("\n".join(self.svg), encoding="utf-8")


def build_business_diagram() -> None:
    c = DiagramCanvas(1800, 1160, "GEO OS 产品业务架构", "从租户与项目边界出发，贯通需求、真实观测、决议、快照、行动与交付")

    c.text(70, 158, "用户与经营边界", font="section", fill="#2E74B5")
    actors = [
        (70, "项目唯一负责人", ["业务/产品/技术/Gate Owner"]),
        (470, "内部交付团队", ["建档、监测、复核、交付"]),
        (870, "受控第三方 Tenant", ["管理本机构成员与自有客户"]),
        (1270, "客户/品牌方", ["项目级只读与交付查看"]),
    ]
    for x, title, lines in actors:
        c.box(x, 196, 330, 104, title, lines, fill="#FFFFFF", outline="#93C5FD")

    c.text(70, 332, "端到端产品价值链", font="section", fill="#167D8D")
    chain = [
        (70, "业务建档", ["Tenant", "Customer / Brand / Project"]),
        (360, "需求与事实", ["Demand & Question", "Brand Truth & Evidence"]),
        (650, "监测与执行", ["Monitoring / Sample Slot", "Execution Run"]),
        (940, "观测与决议", ["Raw Observation", "Assessment / Resolution"]),
        (1230, "测量与洞察", ["Metric Contribution", "Snapshot / Intelligence"]),
        (1520, "行动与交付", ["Intervention / Validation", "Report & Delivery"]),
    ]
    for x, title, lines in chain:
        c.box(x, 372, 240, 146, title, lines, fill="#E7F4F5", outline="#5FB3BE", title_color="#0F5F6B")
    for index in range(len(chain) - 1):
        c.arrow(chain[index][0] + 240, 445, chain[index + 1][0] - 8, 445, fill="#167D8D")

    c.text(70, 558, "产品能力域", font="section", fill="#2F6B45")
    domains = [
        (70, 598, "GEO Control Plane", ["IAM & Tenant", "客户/品牌/项目", "任务、权限、审计"]),
        (390, 598, "Demand & Truth", ["需求信号与问题版本", "品牌事实与证据", "行业规则注入"]),
        (710, 598, "Observation", ["采样与真实端执行", "捕获、质量、Correction", "事实不可变"]),
        (1030, 598, "Measurement", ["Assessment / Review", "Resolution", "两层 Snapshot Membership"]),
        (1350, 598, "Intelligence & Delivery", ["Citation / Source", "策略与干预", "报告只消费 Snapshot"]),
    ]
    for x, y, title, lines in domains:
        c.box(x, y, 280, 154, title, lines, fill="#E8F3EC", outline="#86B89A", title_color="#2F6B45")

    c.text(70, 792, "语义与历史可信层", font="section", fill="#9A6700")
    semantics = [
        (70, "Pack A", ["Observation / Validity / Eligibility", "Review / Snapshot"]),
        (500, "Pack B", ["Mention ≠ Consideration", "≠ Recommendation；贡献三态"]),
        (930, "Pack C", ["Citation / Source / Qualification", "Coverage Context"]),
        (1360, "事实分层", ["Fact → Assessment → Resolution", "→ Snapshot"]),
    ]
    for x, title, lines in semantics:
        c.box(x, 830, 370, 112, title, lines, fill="#FFF4CE", outline="#E5BF62", title_color="#7A5A00")

    c.text(70, 982, "产品分期边界", font="section", fill="#9B1C1C")
    c.box(70, 1020, 505, 92, "Walking Skeleton", ["先验证 Tenant → Observation → Snapshot → Basic Report"], fill="#FDECEC", outline="#E6A2A2", title_color="#9B1C1C")
    c.box(645, 1020, 505, 92, "最小语义扩展", ["Pack B Mention/Recommendation + Pack C Citation Presence"], fill="#FDECEC", outline="#E6A2A2", title_color="#9B1C1C")
    c.box(1220, 1020, 510, 92, "Commercial Complete", ["Coverage Context 后开放正式 KPI，再扩干预与效果验证"], fill="#FDECEC", outline="#E6A2A2", title_color="#9B1C1C")
    c.save(BUSINESS_PNG, BUSINESS_SVG)


def build_technical_diagram() -> None:
    c = DiagramCanvas(1800, 1260, "GEO OS 技术架构", "模块化单体承载核心业务语义，独立执行与异步处理单元通过可重试契约协同")

    c.text(70, 158, "访问层", font="section", fill="#2E74B5")
    access = [
        (70, "GEO Web Console", ["运营配置 / 复核 / 报告"]),
        (470, "Tenant Workspace", ["受控第三方管理自有客户"]),
        (870, "Client Read-only", ["项目级交付查看"]),
        (1270, "Platform Operations", ["平台管理 / 审计 / 排障"]),
    ]
    for x, title, lines in access:
        c.box(x, 198, 330, 102, title, lines, fill="#FFFFFF", outline="#93C5FD")

    c.text(70, 334, "核心应用层", font="section", fill="#167D8D")
    c.box(70, 372, 1180, 316, "GEO Core API · Modular Monolith", [], fill="#E7F4F5", outline="#5FB3BE", title_color="#0F5F6B")
    core_boxes = [
        (100, 424, "IAM & Tenant", ["租户上下文", "RBAC / Audit"]),
        (360, 424, "Customer / Project", ["客户、品牌、项目", "QuestionVersion"]),
        (620, 424, "Contract / Policy", ["Immutable Release", "实际版本绑定"]),
        (880, 424, "Monitoring", ["SamplingDesign", "SampleSlot"]),
        (100, 554, "Observation", ["Finalization", "Correction"]),
        (360, 554, "Decisioning", ["Assessment / Review", "Resolution"]),
        (620, 554, "Measurement", ["MetricContribution", "Snapshot"]),
        (880, 554, "Intelligence / Delivery", ["Citation / Source", "Report / Intervention"]),
    ]
    for x, y, title, lines in core_boxes:
        c.box(x, y, 230, 104, title, lines, fill="#FFFFFF", outline="#8FC9D0", title_color="#0F5F6B")

    c.box(1310, 372, 420, 144, "Contract & Assurance", ["OpenAPI / AsyncAPI / JSON Schema", "ADR / Migration / Contract Test"], fill="#FFF4CE", outline="#E5BF62", title_color="#7A5A00")
    c.box(1310, 544, 420, 144, "Cross-cutting Controls", ["Tenant Context / Policy Release", "Trace ID / Audit / Security"], fill="#FFF4CE", outline="#E5BF62", title_color="#7A5A00")

    for x in (235, 635, 1035, 1475):
        c.arrow(x, 300, x, 360, fill="#2E74B5")

    c.text(70, 728, "独立运行单元与异步协同", font="section", fill="#2F6B45")
    c.box(70, 770, 380, 162, "Real AI Query Engine", ["真实 AI Surface 执行", "响应检测 / 捕获 / 执行证据", "失败不自动形成 Observation"], fill="#E8F3EC", outline="#86B89A", title_color="#2F6B45")
    c.box(510, 770, 380, 162, "AI / Data Worker", ["抽取、分类、Assessment 辅助", "Measurement / Report 任务", "不得覆盖 RawObservation"], fill="#E8F3EC", outline="#86B89A", title_color="#2F6B45")
    c.box(950, 770, 340, 162, "Model Gateway", ["模型与 Prompt 版本", "成本、调用与证据追踪", "不替代真实端观测"], fill="#E8F3EC", outline="#86B89A", title_color="#2F6B45")
    c.box(1350, 770, 380, 162, "Publisher Worker · Deferred", ["发布适配器与人工辅助", "不作为 V1 成立门禁", "后续经范围变更启用"], fill="#FDECEC", outline="#E6A2A2", title_color="#9B1C1C")

    c.arrow(260, 688, 260, 758, fill="#167D8D")
    c.arrow(700, 688, 700, 758, fill="#167D8D")
    c.arrow(1120, 758, 1120, 688, fill="#167D8D")
    c.arrow(440, 850, 500, 850, fill="#2F6B45")
    c.arrow(890, 850, 940, 850, fill="#2F6B45")

    c.text(70, 970, "数据与基础设施层", font="section", fill="#9A6700")
    c.box(70, 1010, 500, 150, "PostgreSQL · Truth Store", ["业务事实、状态机、历史版本", "Outbox / Inbox / Resolution", "Snapshot Membership 与 Provenance"], fill="#FFF4CE", outline="#E5BF62", title_color="#7A5A00")
    c.box(650, 1010, 500, 150, "Object Storage · Evidence", ["截图、原始捕获、文件与清单", "tenant/context 路径与 Hash", "数据库保存不可丢失引用"], fill="#FFF4CE", outline="#E5BF62", title_color="#7A5A00")
    c.box(1230, 1010, 500, 150, "Redis / BullMQ · Runtime Only", ["队列、租约、锁和缓存", "可重建运行时状态", "禁止作为唯一业务事实源"], fill="#FFF4CE", outline="#E5BF62", title_color="#7A5A00")

    c.arrow(300, 932, 300, 998, fill="#9A6700")
    c.arrow(700, 932, 875, 998, fill="#9A6700")
    c.arrow(700, 932, 1480, 998, fill="#9A6700")

    c.rect(70, 1192, 1660, 42, fill="#0B2545", outline="#0B2545", radius=10)
    c.text(900, 1200, "一致性主线：业务事务 + Outbox → 可重复投递 → Inbox/幂等消费 → 领域状态恢复；Report 只读取冻结 Snapshot", font="small", fill="#FFFFFF", anchor="ma")
    c.save(TECH_PNG, TECH_SVG)


def add_document_control(doc: Document) -> None:
    doc.add_heading("1. 修订说明与权威边界", level=1)
    add_callout(
        doc,
        "本次修订结论",
        "现行业务规划和系统逻辑由项目唯一负责人确认。V1.1 不重新发明产品方向，而是把 Gate 0、Walking Skeleton、Tenant 边界、策略版本和正式 KPI 分期写入产品基线。",
    )
    table = doc.add_table(rows=1, cols=3)
    for index, value in enumerate(("调整主题", "V1.1 对齐口径", "对产品的影响")):
        table.cell(0, index).text = value
    rows = [
        ("推进方式", "Gate 0 + Walking Skeleton", "允许基础设施并行，结构性契约分 Gate 关闭"),
        ("Tenant", "受控第三方 Workspace 属于 V1", "自助注册、代理层级、白标、计费、开放 API 后移"),
        ("Policy", "版本化 Manifest + Immutable Release + Code Evaluator", "V1 不建设动态低代码规则平台"),
        ("行业扩展", "首行业 + 影子行业验证无硬编码", "完整 IndustryPolicyPack 在第二行业前冻结"),
        ("客户指标", "Walking Skeleton 不输出正式客户 KPI", "Pack B/C 与 Coverage Context 满足后开放"),
        ("治理", "单负责人兼任多角色", "替补负责人 N/A，保留版本、哈希、批准时间和变更记录"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table, [1900, 4000, 3460], font_size=8.8)


def add_product_positioning(doc: Document) -> None:
    doc.add_heading("2. 产品定位与最高原则", level=1)
    p = doc.add_paragraph()
    p.add_run("GEO OS 是一套围绕“AI 如何参与消费者决策”构建的决策基础设施。")
    p.add_run("它连接真实需求、品牌事实、AI 端面观测、可信测量、诊断、干预和效果验证；内容生产只是干预手段之一。")
    add_bullets(
        doc,
        [
            "Execution（执行）≠ Observation（观测）≠ Measurement（测量）。",
            "Fact（事实）≠ Assessment（评估）≠ Resolution（决议）≠ Snapshot（快照）。",
            "RawObservation 定稿后不可原地修改；错误通过 Correction、Review 和替代链处理。",
            "Validity、Metric Eligibility、Production Inclusion 与 Comparability 分离。",
            "Mention ≠ Consideration ≠ Recommendation；Citation 与 Source 按 Pack C 分层。",
            "任何正式指标必须同时展示 Coverage Context，并绑定实际策略、决议和快照版本。",
        ],
    )
    add_callout(doc, "产品中心", "真实、可追踪、可复现的 AI 决策闭环，而不是监控看板、内容工厂或通用代理商 SaaS。", fill=LIGHT_BLUE, accent=BLUE)


def add_users_and_tenant(doc: Document) -> None:
    doc.add_heading("3. 用户、租户与使用场景", level=1)
    table = doc.add_table(rows=1, cols=3)
    for index, value in enumerate(("用户/边界", "V1 使用场景", "明确限制")):
        table.cell(0, index).text = value
    rows = [
        ("项目唯一负责人", "统一承担业务、产品、技术和 Gate 决策", "每次批准仍需记录版本、哈希与时间"),
        ("内部交付人员", "建档、采样、复核、诊断、报告和干预记录", "不得绕过 Snapshot 或人工改写事实"),
        ("受控第三方 Tenant", "平台开通后管理本机构成员和自有客户", "无自助注册、平台代理层级、白标、分佣和计费"),
        ("客户/品牌方", "在项目授权范围内查看交付和历史快照", "不得访问全局身份目录或其他租户上下文"),
        ("平台运维", "维护端面、任务、账号、审计与排障", "跨租户访问必须带 actor、reason、target 和审计证据"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table, [2100, 3900, 3360], font_size=8.8)

    doc.add_heading("3.1 Ownership Scope", level=2)
    add_bullets(
        doc,
        [
            "SYSTEM_GLOBAL：Platform、Surface、Capability、Contract Definition、Policy Release。",
            "TENANT_OWNED：Customer、Brand、Project、Observation、Snapshot、Report。",
            "GLOBAL_IDENTITY_WITH_TENANT_CONTEXT：SourceDomain、SourceDocument 等可全局去重身份；证据、关系、观测和可见性始终保留租户/项目上下文。",
        ],
    )


def add_business_architecture(doc: Document) -> None:
    doc.add_heading("4. 产品业务架构", level=1)
    p = doc.add_paragraph("业务架构以客户决策闭环为主线，用语义与历史可信层约束各产品能力域。")
    p.paragraph_format.keep_with_next = True
    shape = doc.add_picture(str(BUSINESS_PNG), width=Inches(6.75))
    set_picture_alt(shape, "GEO OS 产品业务架构", "展示用户边界、端到端产品价值链、能力域、A/B/C 语义层以及 Walking Skeleton 到 Commercial Complete 的分期。")
    add_caption(doc, "图 1  GEO OS 产品业务架构 V1.1")

    doc.add_heading("4.1 业务闭环", level=2)
    add_numbers(
        doc,
        [
            "在 Tenant 内建立 Customer、Brand、Project 与默认 Policy Binding。",
            "形成不可变 QuestionVersion 与 SampleSlot；ExecutionRun 可失败和重试，但不增加样本数。",
            "只有用户可见响应结果才进入 ObservationCandidate，并经 Finalization 形成不可变 RawObservation。",
            "Assessment、Review 和 Resolution 追加产生，MetricContribution 明确 VALUE_1、VALUE_0 或 NO_CONTRIBUTION。",
            "MeasurementSnapshot 冻结两层 Membership、实际 Resolution 与策略版本。",
            "Basic Report 只读取指定 Snapshot；洞察、干预和复测继续产生新事实和替代快照。",
        ],
    )


def add_capability_scope(doc: Document) -> None:
    doc.add_heading("5. 产品能力域与范围", level=1)
    table = doc.add_table(rows=1, cols=4)
    for index, value in enumerate(("能力域", "主要对象/能力", "当前阶段", "关键边界")):
        table.cell(0, index).text = value
    rows = [
        ("Control Plane", "IAM、Tenant、Customer、Brand、Project、Task、Audit", "Foundation / Walking Skeleton", "租户边界贯穿所有业务对象"),
        ("Demand & Truth", "DemandSignal、QuestionVersion、Evidence、FactClaim", "最小能力 IN", "完整需求智能和知识治理逐步扩展"),
        ("Monitoring & Execution", "SamplingDesign、SampleSlot、ExecutionRun", "Walking Skeleton IN", "失败不制造观测；重试不增加 Sample N"),
        ("Observation & Decisioning", "RawObservation、Assessment、Review、Resolution", "Gate 0 + Walking Skeleton", "事实不可变；禁止 global latest wins"),
        ("Measurement & Snapshot", "MetricContribution、Membership、Snapshot", "Gate 0 + Walking Skeleton", "Report 只消费冻结 Snapshot"),
        ("Citation & Source", "Occurrence、Qualification、SourceBinding、Resolution", "Pack C 最小能力", "全局身份不直接向租户暴露"),
        ("GEO Intelligence", "GapSymptom、Hypothesis、StrategyAction", "DEFERRED", "先消费 Snapshot；不阻塞主链"),
        ("Intervention & Validation", "Intervention、Window、ImpactEvidence", "DEFERRED", "不在 V1 宣称严格因果"),
        ("Delivery", "Basic Report、周期/复测报告、Delivery Record", "Walking Skeleton IN", "正式客户 KPI 受语义准入控制"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table, [1900, 3000, 1800, 2660], font_size=8.1)


def add_walking_skeleton(doc: Document) -> None:
    doc.add_heading("6. Walking Skeleton 与语义进入顺序", level=1)
    add_callout(doc, "第一条主链", "Tenant → Project → QuestionVersion → SampleSlot → ExecutionRun → RawObservation Finalization → Quality Assessment → Resolution → MetricContribution → MeasurementSnapshot → Basic Report", fill=LIGHT_GREEN, accent=GREEN)
    table = doc.add_table(rows=1, cols=3)
    for index, value in enumerate(("阶段", "加入能力", "完成证明")):
        table.cell(0, index).text = value
    rows = [
        ("Skeleton", "非客户测试指标或最小 ANSWER_OUTCOME", "租户隔离、不可变事实、幂等、快照和报告消费链"),
        ("Pack B", "Mention、Recommendation 与目标集合", "信号、目标、问题适用性和贡献三态可复现"),
        ("Pack C", "Citation Presence、Qualification、Source Binding", "引用存在、来源解析和 Coverage Context 可追溯"),
        ("Customer Metric", "最小正式指标与覆盖展示", "客户报告满足资格、分母、策略和快照准入"),
        ("Expansion", "Demand、Brand Truth、Source Intelligence、Intervention", "形成完整商业闭环和效果验证"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table, [1500, 3400, 4460], font_size=8.7)


def add_semantics(doc: Document) -> None:
    doc.add_heading("7. 核心产品对象与可信语义", level=1)
    table = doc.add_table(rows=1, cols=4)
    for index, value in enumerate(("层", "回答的问题", "主要对象", "不可违反的规则")):
        table.cell(0, index).text = value
    rows = [
        ("Fact", "真实发生了什么", "ExecutionRun、RawObservation、Evidence", "原始事实定稿后不可覆盖"),
        ("Assessment", "按何种方法如何判断", "Quality/Eligibility/Mention/Citation Assessment", "追加、版本化、携带实际策略和证据"),
        ("Resolution", "当前业务语义如何生效", "Observation/Eligibility/Review Resolution", "按 scope + target + context 解析，禁止 latest wins"),
        ("Snapshot", "正式交付时冻结了什么", "Observation/Metric Membership、Provenance", "历史不自动重算，错误通过 Replacement Snapshot"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table, [1300, 2400, 2900, 2760], font_size=8.6)

    doc.add_heading("7.1 Policy 与行业扩展", level=2)
    add_bullets(
        doc,
        [
            "P0 采用 Versioned Policy Manifest + Immutable Policy Release + Code-based Evaluator + Project/Snapshot Binding。",
            "Project Binding 只提供默认值；Execution、Assessment、Resolution 和 Snapshot 保存实际 policy_release_id。",
            "核心模型、状态机、事件和 API 不出现行业硬编码；首行业规则通过版本化配置或接口注入。",
            "完整 IndustryPolicyPack 的依赖、兼容、发布、迁移、停用和回滚在第二行业正式接入前冻结。",
        ],
    )


def add_technical_architecture(doc: Document) -> None:
    doc.add_heading("8. 技术架构", level=1)
    p = doc.add_paragraph("技术架构维持模块化单体与独立运行单元的克制组合，所有业务真相和历史复现以 PostgreSQL 契约为准。")
    p.paragraph_format.keep_with_next = True
    shape = doc.add_picture(str(TECH_PNG), width=Inches(6.75))
    set_picture_alt(shape, "GEO OS 技术架构", "展示访问层、GEO Core 模块化单体、Query Engine、AI/Data Worker、Model Gateway、延后 Publisher Worker，以及 PostgreSQL、对象存储和 Redis 的职责边界。")
    add_caption(doc, "图 2  GEO OS 技术架构 V1.1")

    doc.add_heading("8.1 运行单元与存储职责", level=2)
    table = doc.add_table(rows=1, cols=3)
    for index, value in enumerate(("组件", "正式职责", "禁止事项")):
        table.cell(0, index).text = value
    rows = [
        ("GEO Core API", "租户、项目、策略、观测、决议、快照和交付语义", "不得直接控制真实端面或绕过领域规则改历史"),
        ("Real AI Query Engine", "真实端面执行、响应检测、捕获和执行证据", "不得决定正式 Mention/Citation/Recommendation"),
        ("AI/Data Worker", "异步抽取、评估辅助、测量和报告任务", "不得覆盖 RawObservation 或以当前配置重写历史"),
        ("PostgreSQL", "业务事实、Outbox/Inbox、Resolution、Membership", "不得把唯一事实放在 Redis 或队列"),
        ("Redis/BullMQ", "队列、租约、锁、缓存和运行时调度", "不得作为事实源或历史复现依据"),
        ("Object Storage", "截图、原始捕获、文件和完整性清单", "不得存在无数据库引用的不可追踪正式证据"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table, [2100, 3900, 3360], font_size=8.6)

    doc.add_heading("8.2 异步一致性", level=2)
    add_bullets(
        doc,
        [
            "PostgreSQL 事务内同时写业务状态和 Outbox；投递失败可重试。",
            "消费者使用 message_id + handler_version 或业务幂等键防止重复事实。",
            "Execution 重试产生新的 ExecutionRun，但继续服务同一 SampleSlot。",
            "Worker 租约、Heartbeat、乱序和失败恢复由数据库状态机与可审计失败状态控制。",
        ],
    )


def add_scope_and_roadmap(doc: Document) -> None:
    doc.add_heading("9. Commercial MVP 范围与路线", level=1)
    table = doc.add_table(rows=1, cols=4)
    for index, value in enumerate(("范围组", "IN", "DEFERRED", "OUT / Post-MVP")):
        table.cell(0, index).text = value
    rows = [
        ("客户与行业", "1-3 客户；1-2 Anchor + 1 Shadow", "完整多行业 Policy Pack", "大规模租户与 100 行业基线"),
        ("AI 端面", "2-3 Web；1-2 真机 PoC", "更多端面与移动端生产化", "9 Web + 4 Mobile 一次性全量"),
        ("产品主链", "建档、采样、观测、决议、快照、Basic Report", "Intelligence、Intervention、Effect Validation 完整能力", "渠道大规模自动发布"),
        ("商业平台", "受控第三方 Tenant Workspace", "后续 Agency 产品能力", "自助注册、代理层级、白标、计费、开放 API"),
        ("规则平台", "版本化 Manifest 与代码 Evaluator", "完整 IndustryPolicyPack", "动态低代码规则平台"),
        ("正式 KPI", "Skeleton 仅测试指标", "Pack B/C + Coverage Context 后开放", "无覆盖背景的客户指标"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table, [1450, 2800, 2800, 2310], font_size=8.1)


def add_gate_and_governance(doc: Document) -> None:
    doc.add_heading("10. Gate 0 与单负责人治理", level=1)
    add_callout(
        doc,
        "单负责人治理例外",
        "当前项目由唯一负责人统一承担 Sponsor、PM、Product Owner、Architecture Approver 与 Gate Owner。Backup Owner 标记为 N/A，不虚构多人分工；每次冻结仍必须记录负责人身份、批准时间、精确版本、哈希和已接受风险。",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )
    table = doc.add_table(rows=1, cols=3)
    for index, value in enumerate(("Gate", "必须关闭的结构性决策", "当前状态")):
        table.cell(0, index).text = value
    rows = [
        ("G0-01", "Scope、文档权威、变更控制与单负责人记录", "OPEN / PENDING DOCUMENTATION COMPLETION"),
        ("G0-02", "Tenant、Identity、ownership scope、访问边界", "OPEN；仅允许非约束性准备"),
        ("G0-03", "Observation CAPTURING/FINALIZING/FINALIZED 与 Correction", "OPEN"),
        ("G0-04", "Assessment/Review/Resolution Envelope 与解析协议", "OPEN"),
        ("G0-05", "Outbox/Inbox、幂等、重试、重复和恢复", "OPEN"),
        ("G0-06", "两层 Membership、Resolution ID、替代快照和历史完整性", "OPEN"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table, [1300, 5260, 2800], font_size=8.6)

    doc.add_heading("10.1 允许并行与禁止事项", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "允许并行启动"
    table.cell(0, 1).text = "Gate 未关闭前禁止"
    left = "Monorepo、CI、Lint、测试、Docker Compose、PostgreSQL/Redis/对象存储、迁移框架、认证骨架、Audit/Outbox、Trace ID 与非争议 Skeleton 骨架。"
    right = "冻结核心业务表；实现正式 KPI；各模块自行定义 Tenant、Assessment、Resolution、Outbox 或 Snapshot；用项目当前配置重写历史。"
    table.add_row().cells[0].text = left
    table.rows[1].cells[1].text = right
    style_table(table, [4680, 4680], font_size=8.9)


def add_acceptance(doc: Document) -> None:
    doc.add_heading("11. 产品级验收", level=1)
    add_bullets(
        doc,
        [
            "Tenant A 无法读取 Tenant B 的项目、观测、快照、文件、缓存命中或任务存在性。",
            "执行失败不制造 Observation；重试不增加 Sample N；Finalization 重试不重复生成 RawObservation。",
            "FINALIZED 原始事实不可修改，新 Assessment/Resolution 不覆盖历史。",
            "Project Policy 更新不改变历史 Execution、Assessment、Resolution、Snapshot 和 Report。",
            "Outbox 重投和 Worker 重复消费不产生重复事实，失败可以恢复到确定状态。",
            "Basic Report 重复打开得到同一 Snapshot 结果；正式指标同时展示 Coverage Context。",
            "核心代码通过无行业硬编码测试；第二行业接入前完成完整规则包契约。",
        ],
    )
    add_callout(doc, "最终产品定义", "先把一个真实客户的 AI 决策闭环做深、做稳、做可重复，再扩平台、行业、自动化和商业模式。", fill=LIGHT_BLUE, accent=BLUE)


def add_sources(doc: Document) -> None:
    doc.add_heading("附录 A. 权威输入", level=1)
    add_bullets(doc, SOURCE_DOCS)
    p = doc.add_paragraph()
    p.add_run("领域语义继续以 Decision Pack A1/A2/A3/B/C 最终冻结版为正式来源；本文件负责产品层映射，不反向修改冻结语义。")


def build_docx() -> None:
    doc = Document()
    configure_styles(doc)
    setup_page(doc)
    add_title_block(doc)
    add_document_control(doc)
    add_product_positioning(doc)
    add_users_and_tenant(doc)
    add_business_architecture(doc)
    add_capability_scope(doc)
    add_walking_skeleton(doc)
    add_semantics(doc)
    add_technical_architecture(doc)
    add_scope_and_roadmap(doc)
    add_gate_and_governance(doc)
    add_acceptance(doc)
    add_sources(doc)

    props = doc.core_properties
    props.title = "GEO OS 产品方案对齐修订版 V1.1"
    props.subject = "产品定位、范围、业务架构、技术架构、Gate 0 与 Walking Skeleton"
    props.author = "GEO OS Project"
    props.keywords = "GEO OS, Product Architecture, Technical Architecture, Gate 0, Walking Skeleton"
    props.comments = "Derived from four V1.0 baselines and the latest single-owner governance decision."

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(DOCX_OUTPUT)


def main() -> None:
    outputs = [DOCX_OUTPUT, BUSINESS_SVG, BUSINESS_PNG, TECH_SVG, TECH_PNG]
    existing = [path for path in outputs if path.exists()]
    if existing:
        existing_list = "\n".join(f"- {path}" for path in existing)
        raise FileExistsError(
            "Refusing to overwrite existing review artifacts. Set "
            "GEO_OS_PRODUCT_OUTPUT_DIR to a new versioned directory.\n"
            f"Existing files:\n{existing_list}"
        )
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_business_diagram()
    build_technical_diagram()
    build_docx()
    print(DOCX_OUTPUT)
    print(BUSINESS_SVG)
    print(BUSINESS_PNG)
    print(TECH_SVG)
    print(TECH_PNG)


if __name__ == "__main__":
    main()
