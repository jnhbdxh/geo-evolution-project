from __future__ import annotations

import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(Path(__file__).resolve().parent))

from build_arch_reconciliation_v11 import (  # noqa: E402
    AMBER,
    BLUE,
    DARK_BLUE,
    GREEN,
    INK,
    LIGHT_GRAY,
    MUTED,
    RED_LIGHT,
    add_bullet,
    add_callout,
    add_field,
    add_numbered,
    add_para,
    add_table,
    configure_styles,
    h1,
    h2,
    set_run_font,
)


DEFAULT_OUTPUT = ROOT / "outputs" / "g0-01" / "GEO_OS_G0-01_Formal_Review_Runbook_V1.0.docx"
OUTPUT = Path(os.environ.get("GEO_OS_ARTIFACT_OUTPUT", DEFAULT_OUTPUT))


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("GEO OS  |  G0-01 Formal Review Runbook")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("G0-01 Review Working Document  |  Page ")
    set_run_font(run, size=9, color=MUTED)
    add_field(paragraph, "PAGE")
    run = paragraph.add_run(" of ")
    set_run_font(run, size=9, color=MUTED)
    add_field(paragraph, "NUMPAGES")


def add_cover(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(34)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run("GEO OS")
    set_run_font(run, size=12, bold=True, color=BLUE)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run("G0-01 Formal Review Runbook")
    set_run_font(run, size=23, bold=True, color=INK)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run("Scope Review · Open Decisions · Owner Confirmation · ADR Validation · Closure Record")
    set_run_font(run, size=14, color=DARK_BLUE)

    add_callout(
        doc,
        "控制结论",
        "G0-01 保持 OPEN / PENDING REVIEW。本手册用于组织正式评审，不构成批准，不允许把待评审候选件标记为 APPROVED / FROZEN。",
        fill=AMBER,
        accent="7A5A00",
    )

    add_table(
        doc,
        ["字段", "评审基线"],
        [
            ["手册版本", "V1.0 - REVIEW WORKING DOCUMENT"],
            ["评审候选包", "G0-01 Scope Matrix V1.0 + ADR-001 V1.0"],
            ["候选包状态", "OPEN / PENDING REVIEW"],
            ["完整性控制", "Review Candidate Manifest V1.0（SHA-256）"],
            ["评审顺序", "Scope -> Open Decisions -> Named Owners -> ADR Cases -> Approval -> Closure Record"],
            ["版本规则", "任何内容变化均发布新版本；禁止原地覆盖 V1.0"],
        ],
        [2160, 7200],
    )

    add_para(
        doc,
        "本手册中的空白责任人、结论、日期和签署栏必须由正式评审填写。预填角色或建议不得视为本人确认。",
        color=MUTED,
        italic=True,
    )
    doc.add_page_break()


def add_review_control(doc: Document) -> None:
    h1(doc, "1. Review Control")
    add_para(doc, "目标：在不修改待评审候选件的前提下，对 G0-01 六项退出条件形成可追踪、可签署、可复核的证据。")

    h2(doc, "1.1 Required participants")
    add_table(
        doc,
        ["职责", "具名负责人", "确认方式", "状态"],
        [
            ["Review Chair / Project Manager", "", "主持并确认程序完整", "PENDING"],
            ["Product Sponsor / Approver", "", "批准范围与 Gate 结论", "PENDING"],
            ["Product Owner", "", "确认商业 MVP 与验收边界", "PENDING"],
            ["Domain Architect", "", "确认领域语义与冲突处理", "PENDING"],
            ["Technical Architect", "", "确认技术结构及依赖", "PENDING"],
            ["Security / Data Owner", "", "确认租户、身份与数据边界", "PENDING"],
            ["QA Lead", "", "确认验收和 Contract Test 可执行", "PENDING"],
            ["Recorder", "", "记录决议、异议、版本和时间", "PENDING"],
        ],
        [2450, 2450, 3060, 1400],
        compact=True,
    )

    h2(doc, "1.2 Pre-review integrity checks")
    add_table(
        doc,
        ["Check", "Evidence", "Result / Reviewer"],
        [
            ["候选文件与 Manifest 哈希一致", "两个候选件 SHA-256", ""],
            ["候选件未标记 APPROVED / FROZEN", "文档状态与 Gate Summary", ""],
            ["39 项 Scope 行完整", "Scope Matrix 计数与空值检查", ""],
            ["6 个 Open Decision 均可追踪", "OD-G01-001 至 006", ""],
            ["评审参与者已具名", "上表与会议记录", ""],
            ["评审时间、会议链接/地点已登记", "评审记录", ""],
        ],
        [3000, 3560, 2800],
    )

    add_callout(
        doc,
        "停止条件",
        "若哈希不一致、候选件被覆盖或输入版本不明确，应停止评审，先发布新候选版本和新 Manifest。",
        fill=RED_LIGHT,
        accent="9B1C1C",
    )


def add_scope_review(doc: Document) -> None:
    doc.add_page_break()
    h1(doc, "2. Step 1 - Review 39 Scope Classifications")
    add_para(doc, "先验证全量 39 项，再重点审查 2 项 OUT 与 5 项 DEFERRED 是否对 Walking Skeleton、Commercial MVP 或六项 Gate 形成隐藏依赖。")

    add_table(
        doc,
        ["Scope ID", "Current", "Item", "Required dependency decision"],
        [
            ["X-17", "OUT", "Self-service / Agency / White-label / Billing / Open API", "确认受控第三方 Workspace 不依赖这些开放平台能力。"],
            ["X-18", "OUT", "Dynamic Low-code Policy Platform", "确认版本化 Manifest + Code Evaluator 足以支持 V1 主链。"],
            ["MOD-12", "DEFERRED", "GEO Intelligence", "确认是否属于 Commercial MVP 及最小正式策略验收。"],
            ["MOD-13", "DEFERRED", "Intervention & Effect Validation", "确认代码能力、行业策略和效果验证的分期边界。"],
            ["X-19", "DEFERRED", "Formal Customer Mention / Recommendation / Citation KPI", "确认 Basic Report 不依赖正式客户 KPI。"],
            ["X-20", "DEFERRED", "Full Multi-industry Policy Pack", "确认首行业 + 影子行业足以证明无硬编码。"],
            ["RUN-04", "DEFERRED", "Publisher Worker", "确认 Walking Skeleton 与 Commercial MVP 是否需要独立 Worker。"],
        ],
        [1100, 1150, 2860, 4250],
        compact=True,
    )

    h2(doc, "2.1 Dependency test")
    for text in [
        "移除该项后，Walking Skeleton 是否仍可端到端运行并产生可重复 Snapshot 与 Basic Report？",
        "该项是否被其他 IN 项的验收口径、状态机、事件或数据契约隐式引用？",
        "该项是否影响 Tenant/Identity、不可变观测、Resolution、异步幂等或 Snapshot Contract？",
        "延后是否只影响客户正式能力，而不破坏代码骨架与扩展接口？",
        "若发现依赖，是否已经明确转为 IN、拆分最小 IN 子项，或登记阻塞性 Open Decision？",
    ]:
        add_bullet(doc, text)

    add_table(
        doc,
        ["Review outcome", "Value"],
        [
            ["39 items reviewed", "YES / NO"],
            ["2 OUT accepted", "YES / NO / CHANGES REQUIRED"],
            ["5 DEFERRED accepted", "YES / NO / CHANGES REQUIRED"],
            ["Hidden main-chain dependency found", "NONE / Scope ID(s):"],
            ["Required new candidate version", "NO / YES - target version:"],
            ["Reviewer / Date", ""],
        ],
        [3200, 6160],
    )
    doc.add_page_break()


def add_open_decisions(doc: Document) -> None:
    h1(doc, "3. Step 2 - Dispose Six Open Decisions")
    add_para(doc, "每项必须形成明确决议、所需版本变更和证据；无法决议时继续阻塞 G0-01。不得用‘原则同意’替代可执行结论。")

    add_table(
        doc,
        ["ID", "Decision", "Required evidence", "Review disposition"],
        [
            ["OD-G01-001", "为 39 项指定具名负责人和替补", "本人确认记录", "RESOLVED / BLOCKING"],
            ["OD-G01-002", "GEO Intelligence 是否属于 Commercial MVP", "范围 + 最小验收更新", "RESOLVED / BLOCKING"],
            ["OD-G01-003", "Intervention / Effect Validation 边界", "版本化验收契约", "RESOLVED / BLOCKING"],
            ["OD-G01-004", "Publisher Worker 为 DEFERRED 或 MVP-2", "Runtime scope ADR", "RESOLVED / BLOCKING"],
            ["OD-G01-005", "需求章节到 Scope/Acceptance/Test 覆盖", "Coverage report", "RESOLVED / BLOCKING"],
            ["OD-G01-006", "正式客户 KPI 准入条件", "KPI entry ADR + B/C traceability", "RESOLVED / BLOCKING"],
        ],
        [1300, 3250, 3050, 1760],
        compact=True,
    )

    h2(doc, "3.1 Per-decision record")
    add_table(
        doc,
        ["Field", "To be completed in review"],
        [
            ["Open Decision ID", ""],
            ["Decision", ""],
            ["Rationale and rejected alternatives", ""],
            ["Affected Scope / ADR / API / Schema / Test", ""],
            ["Change classification", "EDITORIAL / COMPATIBLE / BREAKING / MVP SCOPE"],
            ["Artifact and target version", ""],
            ["Owner / Due date", ""],
            ["Disposition", "RESOLVED / REMAINS BLOCKING"],
            ["Approved by / At", ""],
        ],
        [3100, 6260],
    )

    add_callout(
        doc,
        "Gate rule",
        "任一 Open Decision 保持 BLOCKING，或其决议要求的新候选版本尚未评审，G0-01 必须继续 OPEN。",
        fill=AMBER,
        accent="7A5A00",
    )


def add_owner_confirmation(doc: Document) -> None:
    h1(doc, "4. Step 3 - Named Owner Confirmation")
    add_para(doc, "Scope Matrix 中的 Owner Role 不是关闭证据。每一行必须填写主责人与替补负责人，并由本人确认承担验收、变更响应和未决事项处置责任。")

    h2(doc, "4.1 Confirmation procedure")
    steps = [
        "由 Project Manager 将 39 项按 Owner Role 分组并指定具名主责人与替补。",
        "主责人核对 Scope、Acceptance、Authority Source、Gate Dependency 与 Known Limitation。",
        "主责人在评审记录中确认接受；委托确认或仅抄送不计为本人确认。",
        "冲突或无人承接的条目登记为 OD-G01-001 子项，并阻塞关闭。",
        "具名结果写入新版本 Scope Matrix，不覆盖 V1.0。",
    ]
    for step in steps:
        add_numbered(doc, step)

    add_table(
        doc,
        ["Completion evidence", "Requirement", "Result"],
        [
            ["Named primary owner", "39 / 39", ""],
            ["Named backup owner", "39 / 39", ""],
            ["Primary owner confirmation", "39 / 39", ""],
            ["Unowned or disputed items", "0", ""],
            ["Updated Scope Matrix version", "New version if content changed", ""],
            ["Project Manager verification", "Name + timestamp", ""],
        ],
        [3000, 3800, 2560],
    )


def add_adr_validation(doc: Document) -> None:
    h1(doc, "5. Step 4 - Validate ADR-001 Against Real Conflicts")
    add_para(doc, "每个案例都必须能够按‘决策领域 -> 同领域优先级 -> 变更分类 -> 版本/追踪更新’得出唯一处理路径。若规则给出多个合理答案，应修订 ADR 并发布新版本。")

    add_table(
        doc,
        ["Case", "Conflict", "Expected handling"],
        [
            ["RC-01", "受控第三方 Workspace vs Agency / White-label / Open API", "商业方向与 MVP 范围分域处理；V1 只保留受控开通、自有客户管理。"],
            ["RC-02", "全局 Source Identity vs Tenant Isolation", "全局 Canonical Identity 不对租户直露；证据、关系、观测和可见性经 Tenant Context 隔离。"],
            ["RC-03", "策略版本化要求 vs 动态低代码规则平台", "技术结构采用 Manifest + Immutable Release + Code Evaluator；动态平台留待后续 ADR。"],
            ["RC-04", "Project 默认 Policy vs 历史实际 Policy", "Execution、Assessment、Resolution、Snapshot 保存实际 policy_release_id。"],
            ["RC-05", "Basic Report 主链验证 vs 正式客户 KPI", "先验证 Snapshot 消费；正式 KPI 受 Pack B/C 与 OD-G01-006 准入约束。"],
        ],
        [1100, 3700, 4560],
        compact=True,
    )

    h2(doc, "5.1 Pass criteria")
    add_table(
        doc,
        ["Criterion", "Pass condition", "Result / Evidence"],
        [
            ["Decision domain", "每个冲突先被唯一映射到主要决策领域", ""],
            ["Within-domain priority", "冻结、具体、明确替代的新版本规则给出唯一优先级", ""],
            ["Cross-domain interaction", "不允许技术文档静默改写商业范围或领域语义", ""],
            ["Change classification", "变化被归入四类之一并触发正确审批", ""],
            ["Traceability", "Scope、ADR、Contract、Test 与 Release 均有更新路径", ""],
            ["Counterexample", "评审人至少提出一个新冲突案例并获得唯一处理结果", ""],
        ],
        [2500, 4000, 2860],
    )
    doc.add_page_break()


def add_approval_and_closure(doc: Document) -> None:
    h1(doc, "6. Steps 5-6 - Approval and Gate Closure Record")
    add_para(doc, "只有退出条件全部满足、所需新版本已完成复审后，批准人才可形成最终评审结论。Gate Closure Record 必须引用最终批准版本，而不是本手册或待评审 V1.0。")

    h2(doc, "6.1 Final review conclusion")
    add_table(
        doc,
        ["Field", "Final value"],
        [
            ["Review result", "APPROVED / APPROVED WITH NON-BLOCKING LIMITATIONS / REJECTED"],
            ["Blocking Open Decisions", "0 required for closure"],
            ["Approved Scope Matrix version", ""],
            ["Approved ADR-001 version", ""],
            ["Known limitations register", ""],
            ["Approver(s)", ""],
            ["Reviewed at", ""],
            ["Next review / expiry if applicable", ""],
        ],
        [3200, 6160],
    )

    h2(doc, "6.2 Gate Closure Record template")
    add_table(
        doc,
        ["Gate field", "Closure value"],
        [
            ["Gate ID", "G0-01"],
            ["Status", "CLOSED only after final approval"],
            ["Decision / ADR", ""],
            ["Executable Artifact", "Approved Scope Matrix + authority/change-control artifacts"],
            ["Contract Tests", "Coverage / integrity / conflict-case evidence"],
            ["Known Limitations", ""],
            ["Approved Version(s)", ""],
            ["Approved By", ""],
            ["Approved At", ""],
            ["Closure Record ID", ""],
            ["Supersedes", ""],
        ],
        [3000, 6360],
    )

    h2(doc, "6.3 Version publication rule")
    add_bullet(doc, "保留当前 V1.0 候选件和其 SHA-256 Manifest，不做原地覆盖。")
    add_bullet(doc, "评审导致任何内容变化时，发布 V1.1 或更高版本，并生成新的 Manifest。")
    add_bullet(doc, "只有最终批准版本可标记 APPROVED / FROZEN；历史候选版本继续保留其原状态。")
    add_bullet(doc, "Closure Record 引用精确版本和哈希；不得只引用‘最新版’或文件夹。")

    add_callout(
        doc,
        "当前状态",
        "本手册发布后，G0-01 仍为 OPEN / PENDING REVIEW。Gate Closure Record 模板保持空白，不得预签或预填 CLOSED。",
        fill=AMBER,
        accent="7A5A00",
    )


def add_g002_boundary(doc: Document) -> None:
    h1(doc, "7. G0-02 Preparation Boundary")
    add_para(doc, "G0-02 可开展 Tenant 与 Identity 的术语、参与者、对象范围、威胁场景和问题清单准备，但这些产物均为 NON-BINDING。")

    add_table(
        doc,
        ["Allowed preparation", "Prohibited commitment"],
        [
            ["Actor / use-case inventory", "冻结 Tenant / Identity / Membership 核心表"],
            ["Object ownership-scope questions", "冻结 tenant_id 放置规则或 RLS 策略"],
            ["Cross-tenant leakage threat scenarios", "向租户直接暴露全局身份目录"],
            ["Identity lifecycle and audit scenarios", "冻结 API、授权中间件或数据库契约"],
            ["Candidate ADR topics and test cases", "关闭 G0-02 或用草案支持不可逆实现"],
        ],
        [4680, 4680],
    )

    add_para(doc, "G0-02 正式契约工作必须等待 G0-01 的批准 Scope、具名决策人和 Gate Closure Record。", bold_prefix="G0-02 正式契约工作")


def build() -> None:
    if OUTPUT.exists():
        raise FileExistsError(
            f"Refusing to overwrite existing artifact: {OUTPUT}. "
            "Set GEO_OS_ARTIFACT_OUTPUT to a new versioned path."
        )
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_cover(doc)
    add_review_control(doc)
    add_scope_review(doc)
    add_open_decisions(doc)
    add_owner_confirmation(doc)
    add_adr_validation(doc)
    add_approval_and_closure(doc)
    add_g002_boundary(doc)

    properties = doc.core_properties
    properties.title = "GEO OS G0-01 Formal Review Runbook V1.0"
    properties.subject = "Formal review control for G0-01"
    properties.author = "GEO OS Project Governance"
    properties.keywords = "G0-01, Scope, ADR-001, Gate Review, Closure Record"
    properties.comments = "REVIEW WORKING DOCUMENT - not approved or frozen"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
