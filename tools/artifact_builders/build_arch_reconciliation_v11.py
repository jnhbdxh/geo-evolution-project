from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DEFAULT_OUTPUT = ROOT / "docs" / "architecture" / "GEO_OS_Technical_Architecture_Reconciliation_V1.1.docx"
OUTPUT = Path(os.environ.get("GEO_OS_ARTIFACT_OUTPUT", DEFAULT_OUTPUT))

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GREEN = "E8F3EC"
AMBER = "FFF4D6"
RED_LIGHT = "FDECEC"
BLACK = "000000"


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color: str = "D0D5DD", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def style_cell_text(cell, bold: bool = False, color: str = BLACK,
                    size: float = 9.2, center: bool = False) -> None:
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              widths_dxa: list[int], header_fill: str = LIGHT_GRAY,
              compact: bool = False) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, header_fill)
        style_cell_text(cell, bold=True, color=INK, size=9.0, center=idx == 0)
    repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = text
            style_cell_text(cell, size=8.7 if compact else 9.2, center=False)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(3)


def add_field(paragraph, instruction: str, display: str = "1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=11, color=BLACK)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=11, color=BLACK)


def add_para(doc: Document, text: str, bold_prefix: str | None = None,
             color: str = BLACK, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=color, italic=italic)


def add_callout(doc: Document, label: str, text: str, fill: str = BLUE_GRAY,
                accent: str = DARK_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    set_table_geometry(table, [9360])
    set_table_borders(table, color=accent, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.10
    r1 = p.add_run(label + "  ")
    set_run_font(r1, size=10.5, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("GEO OS  |  Technical Architecture Reconciliation V1.1")
    set_run_font(r, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    r = fp.add_run("GEO OS · Gate 0 Controlled Baseline  |  Page ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(fp, "PAGE")
    r = fp.add_run(" of ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(fp, "NUMPAGES")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("GEO OS")
    set_run_font(r, size=12, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Technical Architecture Reconciliation V1.1")
    set_run_font(r, size=25, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("技术架构对齐与 Gate 0 启动基线")
    set_run_font(r, size=15, color=DARK_BLUE)

    add_callout(
        doc,
        "冻结结论",
        "批准启动，受六项 Gate 0 约束。允许脚手架、CI、基础设施和 Walking Skeleton 的非争议部分并行推进；Gate 0 未关闭前，不冻结核心业务表和正式客户指标。",
        fill=GREEN,
        accent="2F6B45",
    )

    add_table(
        doc,
        ["项目", "冻结口径"],
        [
            ["文档版本", "V1.1"],
            ["文档状态", "正式启动基线 / Gate 0 Controlled"],
            ["基准日期", "2026-08-21"],
            ["适用范围", "GEO OS Core Platform V1 与首条 Walking Skeleton"],
            ["权威输入", "总体方案 V1.0、Decision Pack A/B/C、Commercial MVP Scope Freeze 决议"],
            ["批准结论", "APPROVED TO START, SUBJECT TO SIX GATE 0 CONTROLS"],
        ],
        [2160, 7200],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("本文件不是新的业务方法论，也不替代 A/B/C。它负责把已冻结语义转化为可开发、可验证、可追踪的工程边界。")
    set_run_font(r, size=10.5, color=MUTED, italic=True)
    doc.add_page_break()


def h1(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def h2(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 2")


def h3(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 3")


def build() -> None:
    if OUTPUT.exists():
        raise FileExistsError(
            f"Refusing to overwrite existing artifact: {OUTPUT}. "
            "Set GEO_OS_ARTIFACT_OUTPUT to a new versioned path."
        )
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    props = doc.core_properties
    props.title = "GEO OS Technical Architecture Reconciliation V1.1"
    props.subject = "Gate 0 controlled technical architecture baseline"
    props.author = "GEO OS Project"
    props.keywords = "GEO OS, Technical Architecture, Gate 0, Walking Skeleton"

    add_cover(doc)

    h1(doc, "1. 执行摘要与正式决议")
    add_para(doc, "本轮对齐确认现有技术方向继续有效：GEO Core 采用模块化单体；Real AI Query Engine、AI/Data Worker 作为独立运行单元；PostgreSQL 是事实存储；Redis 不得成为业务事实源；对象文件进入 Object Storage。")
    add_para(doc, "本文件解决的不是是否继续完善方法论，而是如何在不丢失整体性的前提下启动工程实现。正式批准状态为“批准启动，受六项 Gate 0 约束”。")
    add_callout(doc, "管理原则", "Gate 0 只控制跨模块、后期难迁移的结构性决策。它不要求所有文档完成后才写代码，也不把未来治理平台提前建设进 V1。")

    h2(doc, "1.1 成功标准")
    for item in [
        "六项 Gate 0 均有 Owner、ADR、可执行产物、契约测试、已知限制和客观状态。",
        "Walking Skeleton 能贯通 Tenant → Execution → Observation → Resolution → Metric → Snapshot → Report。",
        "新增模块复用既有 Tenant、Identity、Policy、Assessment、Outbox 和 Snapshot 协议，不自行定义平行语义。",
        "历史 Snapshot 与 Report 在策略、算法和项目配置变化后仍可准确复现。",
        "V1 第三方机构在受控开通后可管理自有客户，且不存在跨租户数据侧漏。",
    ]:
        add_bullet(doc, item)

    h2(doc, "1.2 非目标")
    for item in [
        "不在 V1 建设低代码动态规则平台。",
        "不在 V1 建设自助注册、代理层级、白标、计费和开放 API。",
        "不在 Walking Skeleton 阶段输出正式客户 Mention、Recommendation 或 Citation KPI。",
        "不要求开工前一次填完全部 Contract Traceability Matrix。",
        "不以微服务数量、消息吞吐量或基础设施复杂度作为架构成熟度指标。",
    ]:
        add_bullet(doc, item)

    h1(doc, "2. 文档权威顺序与变更控制")
    add_table(
        doc,
        ["优先级", "权威来源", "工程解释"],
        [
            ["1", "已批准的冻结决议与勘误", "控制启动边界、Gate 0 和后续变更流程。"],
            ["2", "Decision Pack A/B/C 最终冻结版", "控制 Observation、Metric、Snapshot、Mention、Recommendation、Citation 和 Source 语义。"],
            ["3", "需求/产品/技术/业务背景 V1.0", "控制宏观范围、产品路线、技术栈和非功能要求。"],
            ["4", "本文件及后续 ADR", "将上位语义映射为工程决策，不得反向修改上位业务语义。"],
            ["5", "代码、数据库和 API", "必须证明对上述契约的实现，不得以现有代码惯例覆盖冻结规则。"],
        ],
        [900, 3000, 5460],
    )
    add_para(doc, "冲突处理：下位契约可以在不违反上位语义的前提下细化到子指标；例如 Pack C 可以把 A2 对 CITATION_PARTIAL 的粗粒度默认影响细化到 Citation Presence、Count 和 Source Set，但不得把 UNKNOWN 转换成 0。")
    add_para(doc, "任何需要改变冻结语义的请求必须经过 Change Request → Impact Assessment → Contract Version Upgrade；普通开发 PR 不得完成此类变更。")

    h1(doc, "3. 总体架构基线")
    add_table(
        doc,
        ["运行单元", "核心职责", "禁止事项"],
        [
            ["GEO Web Console", "配置、运营、复核、报告查看", "不得在前端重新计算正式指标或自行解析业务状态。"],
            ["GEO Core API", "租户、客户、项目、问题、监测、干预、权限、策略绑定和审计", "不得直接控制浏览器点击，也不得绕过领域服务修改历史事实。"],
            ["Real AI Query Engine", "资源分配、真实端面执行、响应检测、捕获和执行证据", "不得决定 Mention/Recommendation/Citation 正式指标。"],
            ["AI/Data Worker", "抽取、分类、Assessment、Resolution 辅助、Measurement 和报告任务", "不得覆盖 RawObservation，不得使用项目当前配置重写历史结果。"],
            ["Model Gateway", "内部模型调用、成本和 Prompt 版本记录", "不得替代真实 AI Surface 观测。"],
            ["Publisher Worker", "后续发布适配器与人工辅助流程", "不作为 V1 成立门禁。"],
        ],
        [1800, 3420, 4140],
    )
    add_callout(doc, "事实存储", "PostgreSQL 保存业务事实、历史和 Resolution；Object Storage 保存大对象证据；Redis 只保存可重建运行时状态。任何唯一业务事实不得只存在于 Redis 或队列中。")

    h2(doc, "3.1 模块边界")
    add_table(
        doc,
        ["序号", "模块", "直接下游"],
        [
            ["01", "IAM & Tenant", "所有租户业务模块"],
            ["02", "Customer / Brand / Project", "Demand、Knowledge、Monitoring、Report"],
            ["03", "Contract & Policy Release", "Execution、Assessment、Resolution、Snapshot"],
            ["04", "Demand & Question", "Monitoring、Semantic Assessment"],
            ["05", "Brand Truth & Evidence", "Assessment、Intelligence、Content"],
            ["06", "Monitoring & Sampling", "Execution"],
            ["07", "Query Execution", "Observation Capture"],
            ["08", "Observation Capture", "Assessment / Resolution"],
            ["09", "Assessment / Review / Resolution", "Measurement / Snapshot"],
            ["10", "Measurement & Snapshot", "Intelligence、Report、Verification"],
            ["11", "Citation & Source", "Measurement、Source Intelligence"],
            ["12", "GEO Intelligence", "StrategyAction"],
            ["13", "Intervention & Verification", "New Sampling / Snapshot"],
            ["14", "Report & Delivery", "客户交付"],
            ["15", "Resource / Cost / Audit", "管理与经营分析"],
        ],
        [700, 3100, 5560],
        compact=True,
    )
    add_para(doc, "依赖必须保持单向。Report 只能消费被冻结的 Snapshot；Query Engine 只产生执行和捕获事实；Worker 通过领域命令或受控事件写入对应模块，不得跨 Schema 任意更新。")

    h1(doc, "4. 四条冻结注释")
    add_table(
        doc,
        ["编号", "冻结注释", "工程含义"],
        [
            ["FZ-01", "V1 Tenant Workspace 仅支持受控开通和管理自有客户。", "不实现自助注册、代理层级、白标、计费和开放 API。"],
            ["FZ-02", "全局身份目录不得直接对租户暴露。", "证据、关系、观测与可见性通过租户上下文对象隔离，禁止跨租户信息侧漏。"],
            ["FZ-03", "Project Policy Binding 只是默认配置。", "Execution、Assessment、Resolution 和 Snapshot 分别保存实际 policy_release_id。"],
            ["FZ-04", "Walking Skeleton 不提前形成正式客户 KPI。", "MetricContribution 使用测试指标或冻结的最小 ANSWER_OUTCOME；Basic Report 只验证 Snapshot 消费链。"],
        ],
        [850, 3640, 4870],
    )

    h1(doc, "5. Gate 0 治理模型")
    add_para(doc, "Gate 关闭必须基于客观证据，而不是会议口头确认。Gate 状态只允许 OPEN / CLOSED；尚未具备全部证据时保持 OPEN，并通过 Known Limitations 说明当前边界。")
    add_table(
        doc,
        ["字段", "要求"],
        [
            ["Gate ID", "稳定编号，例如 G0-01。"],
            ["Owner", "唯一直接责任人；协作者另列。"],
            ["Decision / ADR", "记录选择、替代方案和理由。"],
            ["Executable Artifact", "ERD、迁移、状态机、Schema、OpenAPI、AsyncAPI 或运行代码。"],
            ["Contract Tests", "能够自动或确定性验证该 Gate 的关键约束。"],
            ["Known Limitations", "明确尚未解决但不阻止当前阶段的事项。"],
            ["Status", "OPEN / CLOSED。"],
            ["Approved At", "关闭批准时间与批准人。"],
        ],
        [1870, 7490],
    )

    h2(doc, "5.1 初始 Gate Register")
    add_table(
        doc,
        ["Gate", "范围", "最小可执行产物", "初始状态"],
        [
            ["G0-01", "Scope 与权威顺序", "Scope Matrix、Authority ADR、变更流程", "OPEN"],
            ["G0-02", "Tenant 与 Identity", "Ownership Matrix、ID Contract、隔离测试", "OPEN"],
            ["G0-03", "Observation 不可变生命周期", "状态机、DDL 约束、Finalization 测试", "OPEN"],
            ["G0-04", "Assessment / Review / Resolution", "Envelope Schema、Resolution 规则、冲突测试", "OPEN"],
            ["G0-05", "异步一致性", "Outbox/Inbox Contract、幂等矩阵、故障测试", "OPEN"],
            ["G0-06", "Snapshot Contract", "两层 Membership、替代流程、复现测试", "OPEN"],
        ],
        [950, 2540, 4690, 1180],
        compact=True,
    )

    doc.add_page_break()
    h1(doc, "6. G0-01｜Scope 与文档权威顺序")
    h2(doc, "6.1 允许并行启动")
    for item in [
        "Monorepo、TypeScript/Python 工程规范、Lint、测试与构建。",
        "Docker Compose、PostgreSQL、Redis、对象存储和本地开发环境。",
        "迁移框架、认证骨架、Audit 基础设施、Outbox 基础设施。",
        "Trace ID、结构化日志、OpenTelemetry 基础封装。",
        "Walking Skeleton 中不依赖未关闭决策的技术骨架。",
    ]:
        add_bullet(doc, item)
    h2(doc, "6.2 Gate 0 未关闭前禁止冻结")
    for item in [
        "核心业务表最终结构和跨模块外键策略。",
        "正式客户 Mention、Recommendation、Citation KPI。",
        "模块自行定义 Tenant、Identity、Assessment、Resolution、Outbox 或 Snapshot 语义。",
        "只保存项目当前 Policy 而不保存实际 policy_release_id 的历史记录。",
        "以 snapshot + result_json 替代成员和 Provenance 的实现。",
    ]:
        add_bullet(doc, item)

    h1(doc, "7. G0-02｜Tenant 与 Identity Contract")
    add_table(
        doc,
        ["Ownership Scope", "适用对象", "租户可见性与约束"],
        [
            ["SYSTEM_GLOBAL", "Platform、Surface、Capability、Contract Definition、Policy Release", "由平台维护；租户只可引用已发布版本。"],
            ["TENANT_OWNED", "Customer、Brand、Project、Question、Observation、Snapshot、Report", "必须通过租户授权访问；原则上不可迁移所有权。"],
            ["GLOBAL_IDENTITY_WITH_TENANT_CONTEXT", "SourceDomain、SourceDocument 等可复用公共身份", "全局目录不直接暴露；租户通过 Evidence、Binding、Relationship、Projection 引用。"],
        ],
        [2200, 3320, 3840],
    )
    h2(doc, "7.1 V1 Tenant Workspace 边界")
    add_para(doc, "V1 支持平台受控创建 Tenant，Tenant Admin 管理本机构成员并创建自有 Customer、Brand 和 Project。客户侧访问使用 ProjectMembership 与 CLIENT_READONLY，不建立无限组织层级。")
    add_para(doc, "V1 不支持第三方自助注册、代理商父子层级、白标、在线计费、分佣和外部 API。上述能力不得反向改变当前 Tenant 数据安全边界。")
    h2(doc, "7.2 Identity 最小规则")
    for item in [
        "业务 ID 使用不可重用的全局唯一标识；显示编号与内部 ID 分离。",
        "任何外部可访问资源都通过授权检查，不以 UUID 难猜作为安全措施。",
        "Tenant-owned 对象创建时固定 tenant_id；普通业务流程不得变更所有权。",
        "唯一约束、幂等键、缓存键、对象存储路径和队列载荷必须声明 tenant scope。",
        "平台管理员访问租户数据必须产生带 actor、reason 和 target 的 Audit Event。",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "待关闭事项", "SourceDocument 全局身份的合并/拆分、隐藏目录授权方式、tenant-context 引用结构和跨租户去重侧漏测试必须在 G0-02 关闭前形成可执行产物。", fill=AMBER, accent="7A5A00")

    h1(doc, "8. G0-03｜Observation 不可变生命周期")
    add_table(
        doc,
        ["状态", "允许操作", "禁止操作 / 退出条件"],
        [
            ["CAPTURING", "追加 CaptureEvent、RawChunk、证据引用", "不得产生正式 RawObservation；响应事实尚未封存。"],
            ["FINALIZING", "验证目标关联、清单、Hash、必需元数据", "不得接受普通捕获追加；失败回到受控恢复流程。"],
            ["FINALIZED", "追加 Assessment、Review、Correction、Projection", "Raw fact 禁止原地 UPDATE；Finalization 结果不可重复生成。"],
        ],
        [1500, 3660, 4200],
    )
    h2(doc, "8.1 Finalization 原子边界")
    for item in [
        "Finalization 必须固定 Capture Manifest、capture_hash、question_version、platform/surface、effective execution context 和文件引用。",
        "数据库写入与对象清单必须具备可验证的一致性；对象上传成功但 Finalization 失败时进入可恢复的 orphan/pending 状态。",
        "同一 ObservationCandidate 的 Finalization 使用幂等键，重复请求返回同一 RawObservation。",
        "已证实元数据错误通过 CorrectionRecord 追加修正；有效读取由 Raw Value + Applicable Correction 产生。",
        "合法删除通过 Tombstone / Audit 流程处理，不允许业务 Hard Delete。",
    ]:
        add_bullet(doc, item)
    h2(doc, "8.2 最小契约测试")
    add_table(
        doc,
        ["测试", "场景", "预期"],
        [
            ["G03-T01", "同一 idempotency key 重复 Finalize", "只产生一个 RawObservation。"],
            ["G03-T02", "Finalized 后修改 raw_answer", "数据库/服务层同时拒绝。"],
            ["G03-T03", "对象上传后数据库事务失败", "可检测、可恢复，不产生半定稿观测。"],
            ["G03-T04", "时区字段证实错误", "保留原值并追加 CorrectionRecord。"],
        ],
        [1200, 3690, 4470],
    )

    h1(doc, "9. G0-04｜Assessment / Review / Resolution 协议")
    add_para(doc, "统一协议用于复用生命周期、证据、版本和审计，不建立万能业务大表。每个领域保持专用 Payload、状态和索引。")
    add_table(
        doc,
        ["通用 Envelope", "最小字段"],
        [
            ["Assessment", "assessment_id、tenant_id、subject_id、scope、target、context_ref、method、policy_release_id、payload_schema_version、evidence_refs、created_at"],
            ["Human Review", "review_id、assessment scope/target/context、decision、reason_code、evidence_refs、reviewer、supersedes、status"],
            ["Resolution", "resolution_id、scope、target、context_ref、applicable assessment/review IDs、policy_release_id、effective outcome、created_at"],
            ["Projection", "current value、source resolution_id、rebuilt_at；仅为缓存，不是真实来源"],
        ],
        [1900, 7460],
        compact=True,
    )
    h2(doc, "9.1 强制解析规则")
    for item in [
        "AssessmentScope + AssessmentTarget + ResolutionContext 共同定义粒度。",
        "先在每个粒度选择 Current Applicable Assessment，再由 Resolution Policy 聚合；禁止全局 latest assessment wins。",
        "Review 必须带授权、Reason Code 和 Primary Evidence；Reviewer Note 不能作为唯一证据。",
        "硬完整性证据可以推翻过去 Review；Review Conflict 未裁决时 Effective Status = SUSPECT。",
        "所有实际执行记录保存 policy_release_id，Project Binding 只提供默认选择。",
    ]:
        add_bullet(doc, item)

    h1(doc, "10. G0-05｜异步一致性与幂等")
    add_callout(doc, "基本决议", "PostgreSQL 事务内写业务状态与 Outbox Event；发布器将事件投递 BullMQ；消费者使用 Inbox/Processed Message 或等效机制实现重复消费安全。Redis 队列不是事实来源。")
    add_table(
        doc,
        ["风险场景", "必须采用的处理"],
        [
            ["数据库提交成功，队列投递失败", "Outbox 保留未发布事件并重试。"],
            ["同一消息重复投递", "Consumer 使用 message_id + handler_version 幂等处理。"],
            ["Worker 成功但回写超时", "以业务幂等键查询既有结果，不重复创建事实对象。"],
            ["Execution 重试", "新的 ExecutionRun；仍服务同一 SampleSlot，不自动增加 Sample N。"],
            ["Finalization 重试", "同一 Candidate + finalization key 只产生一个 RawObservation。"],
            ["Worker 崩溃并失去租约", "Lease + Heartbeat 过期后可重新领取；终态由数据库状态机决定。"],
            ["事件乱序", "聚合版本/期望状态校验；非法转换进入可审计失败状态。"],
        ],
        [3500, 5860],
    )
    h2(doc, "10.1 统一消息上下文")
    add_para(doc, "所有业务事件和任务至少携带 message_id、correlation_id、trace_id、tenant_id、project_id、aggregate_type、aggregate_id、aggregate_version、event_type、schema_version、occurred_at，以及本次实际 policy_release_id（如适用）。")

    h1(doc, "11. G0-06｜Snapshot Contract")
    add_para(doc, "MeasurementSnapshot 不是结果缓存，而是正式历史测量的封存对象。其事实来源是被冻结的成员关系和实际 Resolution Provenance，而不是未来重新运行同名算法。")
    h2(doc, "11.1 正交状态轴")
    add_table(
        doc,
        ["状态轴", "取值", "含义"],
        [
            ["Lifecycle", "DRAFT / FINALIZED / SUPERSEDED", "描述快照是否定稿、是否被新版本替代。"],
            ["Integrity", "CLEAN / UNDER_REVIEW / IMPACTED", "描述历史结果后来是否发现完整性问题。"],
        ],
        [1700, 3300, 4360],
    )
    h2(doc, "11.2 两层 Membership")
    add_table(
        doc,
        ["对象", "回答的问题", "必须冻结"],
        [
            ["SnapshotObservationMembership", "哪些 Observation 构成基础观测集合？", "snapshot、sample_slot、observation、observation_resolution_id"],
            ["SnapshotMetricMembership", "该 Observation 是否为某个 Metric 贡献？", "metric_type、included、reason、eligibility_resolution_id、production_inclusion_decision_id、assessment IDs"],
        ],
        [2450, 2940, 3970],
        compact=True,
    )
    h2(doc, "11.3 历史替代机制")
    add_para(doc, "Finalized Snapshot 后发现错误时，原数字和 Membership 不自动改写；Integrity 标记为 IMPACTED。如需更正，创建 Replacement Snapshot，并通过 supersedes_snapshot_id / replacement_snapshot_id 形成显式替代链。Report 永久绑定其交付时 Snapshot ID。")

    h1(doc, "12. Policy Release 与实际绑定")
    add_para(doc, "V1 采用 Versioned Policy Manifest + Immutable Release Record + Code-based Evaluator，不建设动态低代码规则平台。")
    add_table(
        doc,
        ["对象", "职责"],
        [
            ["PolicyDefinition", "稳定策略身份、作用域和维护责任。"],
            ["PolicyRelease", "不可变版本、代码/配置摘要、Schema 版本、发布时间和状态。"],
            ["ProjectPolicyBinding", "项目默认选择；仅影响未来执行的缺省值。"],
            ["Execution/Assessment/Resolution Binding", "保存本次实际使用的 policy_release_id。"],
            ["Snapshot Binding", "冻结影响分母、成员、聚合和值的实际 Policy Release。"],
        ],
        [2800, 6560],
    )
    add_callout(doc, "历史保护", "Project 当前 Policy 发生变化时，历史 Execution、Assessment、Resolution、Snapshot 和 Report 不得跟随变化。")

    h1(doc, "13. 行业扩展点")
    h2(doc, "13.1 P0：当前必须冻结")
    for item in [
        "通用表、状态机、API 和事件不得出现口腔、汽车、家装等行业专有字段或枚举。",
        "Project 可以引用行业规则版本；执行和快照保存实际使用版本。",
        "问题分类、事实类型、策略规则、干预目录和报告模板通过接口或版本化配置注入。",
        "第一个 Anchor Industry 提供正式规则；Shadow Industry 用于验证核心模型无行业硬编码。",
    ]:
        add_bullet(doc, item)
    h2(doc, "13.2 P1：第二行业接入前冻结")
    add_para(doc, "完整 IndustryPolicyPack 对象、包依赖、兼容性、发布、迁移、停用和回滚机制，在第二个行业进入正式实现前关闭。当前阶段不阻塞基础设施和 Walking Skeleton。")

    h1(doc, "14. Walking Skeleton")
    add_callout(doc, "目标", "尽早证明最危险的执行—观测—解析—测量—快照—报告主链，而不是先完成全部横向模块。")
    add_table(
        doc,
        ["序号", "节点", "最小能力"],
        [
            ["1", "Tenant", "受控开通、成员和隔离上下文"],
            ["2", "Project", "归属、状态和默认 Policy Binding"],
            ["3", "QuestionVersion", "不可变问题版本"],
            ["4", "SampleSlot", "一个槽位最多贡献一个最终观测"],
            ["5", "ExecutionRun", "可失败、可重试、保存实际上下文"],
            ["6", "RawObservation", "CAPTURING → FINALIZING → FINALIZED"],
            ["7", "Quality Assessment", "最小 ANSWER_OUTCOME / 完整性评估"],
            ["8", "Resolution", "实际 Assessment + Policy Release 解析"],
            ["9", "MetricContribution", "非客户正式测试指标或冻结 ANSWER_OUTCOME"],
            ["10", "MeasurementSnapshot", "两层 Membership 和 Provenance"],
            ["11", "Basic Report", "只读取 Snapshot，验证稳定交付链"],
        ],
        [700, 2550, 6110],
        compact=True,
    )
    h2(doc, "14.1 Walking Skeleton 验收")
    for item in [
        "Tenant A 无法读取 Tenant B 的 Project、Observation、Snapshot、文件或任务。",
        "执行失败不会自动创建 ObservationCandidate。",
        "Execution 重试不会增加 Sample N。",
        "RawObservation Finalized 后不可修改。",
        "新 Assessment/Resolution 不覆盖历史记录。",
        "Outbox 重投和 Worker 重复消费不产生重复事实。",
        "Project Policy 更新不改变历史执行和 Snapshot。",
        "Basic Report 重复打开得到同一 Snapshot 结果。",
    ]:
        add_bullet(doc, item)

    h1(doc, "15. 语义能力进入顺序")
    add_numbered(doc, "Walking Skeleton 先使用非客户测试指标或 ANSWER_OUTCOME，验证基础测量链。")
    add_numbered(doc, "加入 Pack B 最小 Mention 与 Recommendation，实现 Target-specific Assessment 和 MetricContribution。")
    add_numbered(doc, "加入 Pack C Citation Presence，先完成 Applicability、Occurrence、Qualification、Completeness 与 Presence。")
    add_numbered(doc, "加入 Coverage Context 和最小客户指标后，才允许形成正式客户 Basic Report。")
    add_numbered(doc, "再扩展 Demand、Brand Truth、Source Intelligence、Intervention 与 Effect Validation。")

    h1(doc, "16. Contract Traceability Matrix 骨架")
    add_table(
        doc,
        ["Contract", "Rule", "Domain Object", "Owner Module", "Artifact", "Test", "阶段"],
        [
            ["A1-D001", "Response Outcome 后才创建 Candidate", "ObservationCandidate", "Query/Observation", "状态机 + API", "CT-A1-01", "Walking Skeleton"],
            ["A2-D001", "Validity ≠ Eligibility", "Quality / Eligibility Assessment", "Assessment", "Schema + Policy", "CT-A2-01", "Pack B/C 前"],
            ["A3-D013", "Snapshot 冻结两层 Membership 与实际决议", "MeasurementSnapshot", "Measurement", "ERD + DDL", "CT-A3-13", "Gate 0"],
            ["B-F02", "VALUE_1 / VALUE_0 / NO_CONTRIBUTION", "MetricContribution", "Measurement", "Schema + Evaluator", "CT-B-F02", "Pack B"],
            ["C-D003", "Occurrence / Qualification / Extraction 分层", "Citation objects", "Citation & Source", "ERD + API", "CT-C-D003", "Pack C"],
            ["FZ-03", "实际 policy_release_id 必须固定", "Execution/Assessment/Resolution/Snapshot", "Cross-cutting", "DDL + Test", "CT-FZ-03", "Gate 0"],
        ],
        [900, 2300, 1800, 1500, 1200, 860, 800],
        compact=True,
    )
    add_para(doc, "矩阵在 Gate 0 建立骨架，随模块开发持续补齐。任何进入正式交付的 Contract Rule 必须能定位到代码、数据库、API/事件和测试。")

    h1(doc, "17. ADR 与可执行产物清单")
    add_table(
        doc,
        ["ADR", "主题", "最低产物", "Gate"],
        [
            ["ADR-001", "文档权威顺序与变更", "Authority Matrix", "G0-01"],
            ["ADR-002", "Tenant Ownership Scopes", "Ownership ERD + 隔离测试", "G0-02"],
            ["ADR-003", "Identity 与不可重用 ID", "ID Contract + 唯一约束", "G0-02"],
            ["ADR-004", "Observation Finalization", "状态机 + DDL + API", "G0-03"],
            ["ADR-005", "Assessment/Review/Resolution", "Envelope Schema + 领域示例", "G0-04"],
            ["ADR-006", "Outbox/Inbox 与幂等", "事件 Schema + 失败矩阵", "G0-05"],
            ["ADR-007", "Snapshot Membership", "ERD + 替代流程", "G0-06"],
            ["ADR-008", "Policy Release Binding", "Manifest Schema + Binding Test", "G0-04/G0-06"],
        ],
        [1200, 3150, 3650, 1360],
        compact=True,
    )

    h1(doc, "18. 下一执行序列")
    add_numbered(doc, "建立三套工程基线目录：Technical Architecture、Core Domain & Data Contracts、Execution Assurance。")
    add_numbered(doc, "创建 Gate Register 和 ADR 模板，指定六项 Gate Owner。")
    add_numbered(doc, "并行启动 Monorepo、CI、Docker Compose、迁移、认证、审计、Outbox 和可观测性骨架。")
    add_numbered(doc, "以 G0-02 至 G0-06 的最小可执行产物驱动 Walking Skeleton。")
    add_numbered(doc, "完成 Walking Skeleton 契约测试后，引入 Pack B/C 最小语义能力。")
    add_numbered(doc, "所有 Gate 关闭证据和 Contract 映射持续写入 Traceability Matrix。")
    add_callout(doc, "启动状态", "本文件批准工程启动，但不宣告任何 Gate 已关闭。Gate 的关闭必须由后续 ADR、Schema、迁移、运行代码和 Contract Test 共同证明。", fill=GREEN, accent="2F6B45")

    h1(doc, "附录 A｜Gate 关闭记录模板")
    add_table(
        doc,
        ["字段", "填写内容"],
        [
            ["Gate ID", ""],
            ["Owner", ""],
            ["Decision / ADR", ""],
            ["Executable Artifact", ""],
            ["Contract Tests", ""],
            ["Known Limitations", ""],
            ["Status", "OPEN"],
            ["Approved At", ""],
        ],
        [2500, 6860],
    )

    h1(doc, "附录 B｜冻结声明")
    add_para(doc, "本文件经确认后作为 GEO OS 后续项目推进的正式管理与技术启动基线。任何团队不得以“先实现再补契约”为理由绕过六项 Gate 0；同样不得以“Gate 未全部关闭”为理由阻止已经明确允许的脚手架、基础设施和 Walking Skeleton 非争议工作。")
    add_para(doc, "正式结论：APPROVED TO START, SUBJECT TO SIX GATE 0 CONTROLS。")

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
